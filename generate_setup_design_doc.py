"""Generate TSG Setup, Design & Implementation Document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ═══════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════
title = doc.add_heading('TMO QA Automation Platform', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Setup, Design & Implementation Guide')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2 = doc.add_paragraph('From Zero to Production-Ready Test Suites')
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ═══════════════════════════════════════════════════════════════
# 1. THE PROBLEM
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. The Problem We Solved', level=1)
doc.add_paragraph(
    'TMO/Spectrum Mobile provisioning has 100+ features across PI-46 to PI-55, '
    'each requiring comprehensive test suites with detailed test steps. '
    'Manual test case writing was taking 2-3 days per feature, with inconsistent quality, '
    'missing coverage, and no traceability back to design specifications.')
doc.add_paragraph('Key challenges:')
for item in [
    'Test cases were written manually from Chalk specs and Jira stories — slow and error-prone',
    'No systematic way to ensure all Chalk scenarios were covered',
    'Steps were generic ("verify the result") instead of domain-specific',
    'No distinction between API, UI, CDR, Inquiry, and Sync features in test steps',
    'Jira subtask tables (e.g., YL state change matrix) were ignored in test generation',
    'No automated execution — all testing was manual',
]:
    doc.add_paragraph(item, style='List Bullet')

# ═══════════════════════════════════════════════════════════════
# 2. WHAT WE BUILT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. What We Built', level=1)
doc.add_paragraph('Three integrated platforms that form a complete QA automation pipeline:')

doc.add_heading('2.1 Test Suite Generator (TSG)', level=2)
doc.add_paragraph(
    'Automatically generates production-ready test suites from Chalk design specs, '
    'Jira requirements, and NBOP UI crawler data. Outputs Excel files with Summary, '
    'Description, Preconditions, Test Steps, and Expected Results for every test case.')

doc.add_heading('2.2 Test Suite Executor (TSE)', level=2)
doc.add_paragraph(
    'Executes the generated test suites automatically. Routes API steps to cURL execution, '
    'UI steps to Playwright browser automation, and Century Report steps to automated validation. '
    'Captures evidence (screenshots, API responses, DOCX reports) per test case.')

doc.add_heading('2.3 MDA Jira Dashboard', level=2)
doc.add_paragraph(
    'Project tracking dashboard that provides visibility into feature status, '
    'test coverage, and execution progress across all PI iterations.')

# ═══════════════════════════════════════════════════════════════
# 3. PREREQUISITES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Prerequisites & Initial Setup', level=1)

doc.add_heading('3.1 Environment Requirements', level=2)
t = doc.add_table(rows=7, cols=2)
t.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate([
    ('Requirement', 'Details'),
    ('Python', '3.11+ (tested on 3.13)'),
    ('Browser', 'Microsoft Edge or Google Chrome'),
    ('Network Access', 'Chalk (chalk.charter.com), Jira (jira.charter.com), NBOP portal, Genesis portal'),
    ('Authentication', 'Corporate SSO credentials for Chalk, Jira, NBOP, Genesis'),
    ('Disk Space', '~500MB for DB cache, outputs, and checkpoints'),
    ('OS', 'Windows 10/11 (tested), macOS/Linux compatible'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_heading('3.2 Initial Setup Steps', level=2)
for i, step in enumerate([
    'Install Python 3.11+ and pip',
    'Clone the repository from GitHub',
    'Install dependencies: pip install -r requirements.txt',
    'Install Playwright browsers: playwright install chromium',
    'Configure .env file with NBOP credentials (optional for auto-login)',
    'Run first-time Chalk sync from the TSG Dashboard to populate the DB cache',
    'Verify by generating a test suite for any PI-53 feature',
], 1):
    doc.add_paragraph('%d. %s' % (i, step))

doc.add_heading('3.3 Key Dependencies', level=2)
for dep in [
    'Streamlit — Web dashboard framework',
    'Playwright — Browser automation for Chalk/Jira/NBOP/Genesis crawling',
    'openpyxl — Excel file generation',
    'python-docx — DOCX evidence document generation',
    'BeautifulSoup — HTML parsing for Chalk content',
    'SQLite — Local database for caching (no external DB required)',
]:
    doc.add_paragraph(dep, style='List Bullet')

# ═══════════════════════════════════════════════════════════════
# 4. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Architecture & Design', level=1)

doc.add_heading('4.1 High-Level Flow', level=2)
doc.add_paragraph(
    'The platform follows a pipeline architecture where each stage enriches the test suite:')
for i, stage in enumerate([
    'Data Collection: Chalk scenarios + Jira (AC, subtasks, comments, attachments) + NBOP UI map',
    'Feature Classification: Detect API / UI / Hybrid / CDR / Inquiry / Sync / Batch',
    'Chalk-to-TC Conversion: Each Chalk scenario becomes a TC with domain-specific steps',
    'Jira Subtask Mining: Parse wiki tables into structured TCs with dynamic steps',
    'Integration Contract: Add MUST CALL / MUST NOT CALL assertions per operation',
    'Test Analyst Reasoning: Gap detection and supplementary TC generation',
    'Scenario Enrichment: Mandatory negatives, audit-history checklist, E2E flows',
    'UI Verification Mirror: NBOP verification TCs for API operations',
    'Quality Gate: Cross-contamination cleanup, precondition validation, category normalization',
    'Humanization: Dedup, priority scoring, description rewriting, final validation',
    'Excel Generation: Production-ready output with formatted sheets',
], 1):
    doc.add_paragraph('%d. %s' % (i, stage))

doc.add_heading('4.2 Key Design Decisions', level=2)

doc.add_heading('Integration Contract (Single Source of Truth)', level=3)
doc.add_paragraph(
    'Instead of hardcoding per-feature logic, we built a global Integration Contract with '
    '24 registered operations and 9 external systems. Every module in the pipeline consults '
    'this contract to determine which systems each operation touches and which it must NOT touch. '
    'This eliminated the "whack-a-mole" problem of fixing one feature and breaking another.')

doc.add_heading('Domain-Specific Step Templates', level=3)
doc.add_paragraph(
    '15+ step templates cover every operation type. Each template generates steps that are '
    'specific to the operation — not generic "verify the result" but "Verify Syniverse '
    'CreateSubscriber is called with correct IMSI, MDN, and wholesale plan." '
    'Templates include: Swap MDN (PSIM vs ESIM), Sync Subscriber (YL/YD/YM/YP/PL), '
    'Order Inquiry, CDR Processing, Hotline (dual assertion), Kafka EVENT_MESSAGES, etc.')

doc.add_heading('Jira Subtask Table Intelligence', level=3)
doc.add_paragraph(
    'When Jira subtasks contain structured wiki tables (e.g., the YL state change matrix '
    'with TMO Status, NSL Status, Syniverse Action columns), the engine parses them into '
    'structured data and generates TCs with steps derived directly from each table row. '
    'This ensures the test suite matches the exact specification in Jira, not just Chalk.')

doc.add_heading('Quality Gate (Multi-Layer)', level=3)
doc.add_paragraph(
    'Every TC passes through a quality gate that checks: trailing dots, empty fields, '
    'non-standard categories, generic expected results, cross-contamination (CDR features '
    'never get API steps, inquiry features never get Century Report steps), and minimum '
    'step count. The gate also normalizes preconditions based on feature type.')

doc.add_heading('4.3 Module Structure', level=2)
t2 = doc.add_table(rows=16, cols=2)
t2.style = 'Light Grid Accent 1'
for i, (mod, desc) in enumerate([
    ('Module', 'Purpose'),
    ('test_engine.py', 'Main pipeline: Chalk-to-TC conversion, Jira mining, enrichment'),
    ('test_analyst.py', 'QA reasoning engine: gap detection, scenario suggestions'),
    ('step_templates.py', '15+ domain-specific step chain templates'),
    ('tc_templates.py', 'Feature classification, description/precondition/step builders'),
    ('integration_contract.py', '24 operations, 9 external systems, Syniverse assertion matrix'),
    ('scenario_enricher.py', 'Mandatory negatives, audit checklist, E2E enrichment'),
    ('ui_mirror.py', 'NBOP UI verification mirror TCs'),
    ('humanizer.py', 'Dedup, priority scoring, description rewriting, final validation'),
    ('chalk_parser.py', 'Chalk page scraping and scenario extraction'),
    ('jira_fetcher.py', 'Jira REST API: issues, subtasks, comments, attachments'),
    ('nbop_ui_knowledge.py', 'NBOP UI knowledge base: menus, fields, navigation paths'),
    ('excel_generator.py', 'Formatted Excel output with multiple sheets'),
    ('database.py', 'SQLite cache for Chalk, Jira, test suites, transactions'),
    ('pipeline.py', 'Orchestration: block-based execution with retry and self-heal'),
]):
    t2.cell(i, 0).text = mod
    t2.cell(i, 1).text = desc

# ═══════════════════════════════════════════════════════════════
# 5. IMPLEMENTATION JOURNEY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Implementation Journey', level=1)

phases = [
    ('Phase 1: Foundation',
     'Built the core pipeline: Chalk parser, Jira fetcher, basic TC generation from scenarios. '
     'Output: Excel files with TC summary and description. Coverage: ~40% of Chalk scenarios.'),
    ('Phase 2: Step Templates',
     'Added domain-specific step templates for each operation type (Activation, Swap MDN, '
     'Change SIM, etc.). Steps went from generic to specific. Coverage: ~70%.'),
    ('Phase 3: Integration Contract',
     'Built the global Integration Contract with 24 operations and 9 external systems. '
     'Added MUST CALL / MUST NOT CALL assertions. Eliminated cross-contamination. Coverage: ~85%.'),
    ('Phase 4: Jira Deep Mining',
     'Added subtask description mining, wiki table parsing with dynamic step generation, '
     'comment mining, and attachment parsing. Coverage: ~95%.'),
    ('Phase 5: Quality Gate & Humanization',
     'Multi-layer quality gate: surface quality, semantic alignment, cross-contamination cleanup, '
     'inquiry/CDR/sync-specific step validation. Humanizer for dedup and priority scoring.'),
    ('Phase 6: UI Mirror & NBOP Knowledge',
     'NBOP UI crawler cache with real menu paths and field names. UI Verification Mirror '
     'auto-generates NBOP verification TCs for API operations.'),
    ('Phase 7: Sync Subscriber & Kafka',
     'Sync Subscriber template with YL/YD/YM/YP/PL state change flows. Kafka EVENT_MESSAGES '
     'template. Dynamic steps from Jira subtask tables.'),
    ('Phase 8: Test Suite Executor (TSE)',
     'Automated execution of generated test suites. API steps via cURL, UI steps via Playwright, '
     'Century Report validation, evidence capture with DOCX reports.'),
]
for title_text, desc in phases:
    doc.add_heading(title_text, level=2)
    doc.add_paragraph(desc)

# ═══════════════════════════════════════════════════════════════
# 6. RESULTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Results', level=1)

t3 = doc.add_table(rows=9, cols=3)
t3.style = 'Light Grid Accent 1'
for i, (metric, before, after) in enumerate([
    ('Metric', 'Before (Manual)', 'After (TSG)'),
    ('Time per feature', '2-3 days', '~30 seconds'),
    ('TC quality consistency', 'Variable', '100% pass rate'),
    ('Chalk coverage', '~50-60%', '97.5%'),
    ('Jira subtask coverage', '~20%', '100% (table rows parsed)'),
    ('Cross-contamination', 'Common', 'Zero'),
    ('Syniverse assertions', 'Often missing', 'Automatic (MUST/MUST NOT)'),
    ('UI verification', 'Separate manual effort', 'Auto-generated mirror TCs'),
    ('Execution', 'Fully manual', 'Automated via TSE'),
]):
    t3.cell(i, 0).text = metric
    t3.cell(i, 1).text = before
    t3.cell(i, 2).text = after

doc.add_paragraph('')
doc.add_paragraph(
    'PI-52 & PI-53 combined: 54 features, 1,789 test cases, 7,233 test steps — '
    'all generated automatically with zero quality issues.')

# ═══════════════════════════════════════════════════════════════
# 7. HOW TO USE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. How to Use (Day-to-Day)', level=1)
for i, step in enumerate([
    'Open TSG Dashboard: cd TestSuiteGenerator && streamlit run TSG_Dashboard_V4.1.py',
    'Select PI iteration (e.g., PI-53)',
    'Select feature from dropdown (e.g., MWTGPROV-4009)',
    'Click "Generate" — suite is built in ~30 seconds',
    'Download the Excel file from the output',
    'Open TSE Dashboard to execute the suite automatically (optional)',
], 1):
    doc.add_paragraph('%d. %s' % (i, step))

out = 'TestSuiteGenerator/outputs/TSG_Setup_Design_Implementation.docx'
doc.save(out)
print('Saved: %s' % out)
