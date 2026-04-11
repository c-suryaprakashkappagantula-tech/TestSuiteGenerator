"""
doc_parser.py — Parse .docx, .xlsx, .pdf files into structured text.
Used for Jira attachments and user-uploaded HLD/LLD docs.
"""
from pathlib import Path
from typing import List, Dict
from dataclasses import dataclass, field


@dataclass
class ParsedDoc:
    filename: str = ''
    file_type: str = ''
    paragraphs: List[str] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)
    open_items: List[str] = field(default_factory=list)
    raw_text: str = ''


def parse_file(filepath: Path, log=print) -> ParsedDoc:
    """Auto-detect file type and parse."""
    fp = Path(filepath)
    doc = ParsedDoc(filename=fp.name, file_type=fp.suffix.lower())

    if doc.file_type == '.docx':
        return _parse_docx(fp, doc, log)
    elif doc.file_type in ('.xlsx', '.xls'):
        return _parse_xlsx(fp, doc, log)
    elif doc.file_type == '.pdf':
        return _parse_pdf(fp, doc, log)
    elif doc.file_type in ('.htm', '.html'):
        return _parse_html(fp, doc, log)
    elif doc.file_type in ('.txt', '.csv'):
        return _parse_text(fp, doc, log)
    else:
        log(f'[DOC] WARN Unsupported file type: {doc.file_type}')
        return doc


# Point 8: Expanded open item keywords
_OPEN_ITEM_KEYWORDS = [
    'cannot ', 'need to ', 'tbd', 'to be determined', 'to be confirmed',
    'pending ', 'open question', 'open item', 'not yet ', 'awaiting ',
    'under discussion', 'to be decided', 'clarification needed',
    'follow up', 'follow-up', 'action item', 'blocker',
]


def _parse_docx(fp, doc: ParsedDoc, log=print) -> ParsedDoc:
    """Parse a .docx file into paragraphs and tables."""
    try:
        from docx import Document
        d = Document(str(fp))

        for p in d.paragraphs:
            t = p.text.strip()
            if t:
                doc.paragraphs.append(t)
                # Point 8: Detect open items with expanded keywords
                t_low = t.lower()
                if any(kw in t_low for kw in _OPEN_ITEM_KEYWORDS):
                    doc.open_items.append(t)

        for table in d.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                doc.tables.append(rows)

        doc.raw_text = '\n'.join(doc.paragraphs)
        log(f'[DOC] OK Parsed {fp.name}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(doc.open_items)} open items')
    except Exception as e:
        log(f'[DOC] FAIL to parse {fp.name}: {e}')
    return doc


def _parse_xlsx(fp, doc: ParsedDoc, log=print) -> ParsedDoc:
    """Parse an .xlsx file — each sheet becomes a table."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(fp), data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(c for c in cells):
                    rows.append(cells)
            if rows:
                doc.tables.append(rows)
                # Also add as text
                for r in rows:
                    doc.paragraphs.append(' | '.join(r))

        doc.raw_text = '\n'.join(doc.paragraphs)
        log(f'[DOC] OK Parsed {fp.name}: {len(doc.tables)} sheets as tables')
    except Exception as e:
        log(f'[DOC] FAIL to parse {fp.name}: {e}')
    return doc


def _parse_pdf(fp, doc: ParsedDoc, log=print) -> ParsedDoc:
    """Parse a .pdf file into text."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(fp))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                for ln in text.split('\n'):
                    ln = ln.strip()
                    if ln:
                        doc.paragraphs.append(ln)

        doc.raw_text = '\n'.join(doc.paragraphs)
        log(f'[DOC] OK Parsed {fp.name}: {len(doc.paragraphs)} lines from {len(reader.pages)} pages')
    except Exception as e:
        log(f'[DOC] FAIL to parse {fp.name}: {e}')
    return doc


def _parse_text(fp, doc: ParsedDoc, log=print) -> ParsedDoc:
    """Parse plain text / CSV."""
    try:
        text = fp.read_text(encoding='utf-8', errors='ignore')
        doc.paragraphs = [ln.strip() for ln in text.split('\n') if ln.strip()]
        doc.raw_text = text
        # Point 8: open item detection for text files too
        for ln in doc.paragraphs:
            ln_low = ln.lower()
            if any(kw in ln_low for kw in _OPEN_ITEM_KEYWORDS):
                doc.open_items.append(ln)
        log(f'[DOC] OK Parsed {fp.name}: {len(doc.paragraphs)} lines, {len(doc.open_items)} open items')
    except Exception as e:
        log(f'[DOC] FAIL to parse {fp.name}: {e}')
    return doc


def _parse_html(fp, doc: ParsedDoc, log=print) -> ParsedDoc:
    """Point 9: Parse .html/.htm files (e.g. Century Report SERVICE_GROUPING exports)."""
    try:
        import re as _re
        text = fp.read_text(encoding='utf-8', errors='ignore')
        # Strip HTML tags to get plain text
        clean = _re.sub(r'<script[^>]*>.*?</script>', '', text, flags=_re.DOTALL | _re.IGNORECASE)
        clean = _re.sub(r'<style[^>]*>.*?</style>', '', clean, flags=_re.DOTALL | _re.IGNORECASE)
        clean = _re.sub(r'<[^>]+>', ' ', clean)
        clean = _re.sub(r'&nbsp;', ' ', clean)
        clean = _re.sub(r'&amp;', '&', clean)
        clean = _re.sub(r'&lt;', '<', clean)
        clean = _re.sub(r'&gt;', '>', clean)
        clean = _re.sub(r'\s+', ' ', clean)

        doc.paragraphs = [ln.strip() for ln in clean.split('\n') if ln.strip() and len(ln.strip()) > 3]
        doc.raw_text = clean

        # Extract HTML tables
        table_pat = _re.compile(r'<table[^>]*>(.*?)</table>', _re.DOTALL | _re.IGNORECASE)
        row_pat = _re.compile(r'<tr[^>]*>(.*?)</tr>', _re.DOTALL | _re.IGNORECASE)
        cell_pat = _re.compile(r'<t[dh][^>]*>(.*?)</t[dh]>', _re.DOTALL | _re.IGNORECASE)

        for tbl_match in table_pat.finditer(text):
            rows = []
            for row_match in row_pat.finditer(tbl_match.group(1)):
                cells = [_re.sub(r'<[^>]+>', '', c).strip() for c in cell_pat.findall(row_match.group(1))]
                if any(cells):
                    rows.append(cells)
            if rows:
                doc.tables.append(rows)

        # Open item detection
        for ln in doc.paragraphs:
            ln_low = ln.lower()
            if any(kw in ln_low for kw in _OPEN_ITEM_KEYWORDS):
                doc.open_items.append(ln)

        log(f'[DOC] OK Parsed {fp.name}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(doc.open_items)} open items')
    except Exception as e:
        log(f'[DOC] FAIL to parse {fp.name}: {e}')
    return doc
