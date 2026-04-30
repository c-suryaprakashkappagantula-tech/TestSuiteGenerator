"""
preload_cabot_chalk.py — Deep-scrape the Cabot Chalk page tree and store in
a dedicated SQLite database (CABOT_CHALK_DB.db) + DOCX report.

Target URL: https://chalk.charter.com/spaces/SPECTRUMMOBIT2/pages/3247714982/Cabot
Space:      SPECTRUMMOBIT2
Page ID:    3247714982

Known child pages (from sidebar):
  - LLD - NMP - Cabot Transfer Lines Lookup
  - Updates to Downstream NSL API calls
  - Update to Create Line API
  - Update to Line Details API
  - Update to Line Summary APIs
  - Data Bolt-On (DBO) Reporting Metrics
  - Get Data Pack Status API Document
  - International Roaming Rates API
  - LineSummary API Documentation

Behavior:
  1. Loads the main Cabot page, scrapes right-panel content (tables, text, links)
  2. Expands the Cabot sidebar tree and discovers ALL child pages
  3. For EACH child page:
     a. Navigates to it, scrapes tables / text / links
     b. Discovers sub-pages (grandchildren) via sidebar tree + in-page child macros
     c. Recursively scrapes every sub-page found (unlimited depth)
  4. Expands any collapsed sections / tabs / expand macros on each page
  5. Stores everything in CABOT_CHALK_DB.db -> Cabot_Chalk table
  6. Generates DOCX: outputs/Cabot_Chalk.docx

Usage:
    python preload_cabot_chalk.py
    python preload_cabot_chalk.py --force            (re-fetch even if DB has data)
    python preload_cabot_chalk.py --headless false    (visible browser for debugging)
    python preload_cabot_chalk.py --max-depth 5       (limit recursion depth, default=10)
"""
import sys, os, time, argparse, json, re, sqlite3
from pathlib import Path
from datetime import datetime

sys.path.insert(0, os.path.dirname(__file__))

from playwright.sync_api import sync_playwright
from modules.config import get_browser_channel, PAGE_LOAD_TIMEOUT_MS, NETWORK_IDLE_TIMEOUT_MS, ROOT

# ================================================================
# CONSTANTS
# ================================================================
CABOT_URL = 'https://chalk.charter.com/spaces/SPECTRUMMOBIT2/pages/3247714982/Cabot'
CABOT_PAGE_ID = '3247714982'
SPACE_KEY = 'SPECTRUMMOBIT2'
DB_PATH = ROOT / 'CABOT_CHALK_DB.db'
DOCX_PATH = ROOT / 'outputs' / 'Cabot_Chalk.docx'
DEFAULT_MAX_DEPTH = 10


def log(msg):
    print(msg, flush=True)


def _wait(page, timeout=NETWORK_IDLE_TIMEOUT_MS):
    try:
        page.wait_for_load_state('networkidle', timeout=timeout)
    except Exception:
        pass


# ================================================================
# DATABASE — standalone CABOT_CHALK_DB.db
# ================================================================

def _conn():
    c = sqlite3.connect(str(DB_PATH), timeout=10)
    c.execute('PRAGMA journal_mode=WAL')
    c.row_factory = sqlite3.Row
    return c


def init_cabot_table():
    """Create the Cabot_Chalk table if it doesn't exist."""
    c = _conn()
    c.executescript('''
        CREATE TABLE IF NOT EXISTS Cabot_Chalk (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            section_name    TEXT NOT NULL,
            section_url     TEXT,
            parent_section  TEXT DEFAULT 'Cabot',
            depth           INTEGER DEFAULT 0,
            page_id         TEXT,
            content_type    TEXT DEFAULT 'table',
            table_data_json TEXT,
            raw_text        TEXT,
            links_json      TEXT,
            sub_sections_json TEXT,
            tab_data_json   TEXT,
            last_fetched    TEXT,
            UNIQUE(section_name, parent_section, section_url)
        );
    ''')
    c.commit()
    c.close()
    log('  [DB] Cabot_Chalk table ready in %s' % DB_PATH)


def save_cabot_section(section_name, section_url, parent_section, depth, page_id,
                       content_type, table_data, raw_text, links, sub_sections,
                       tab_data=None):
    c = _conn()
    now = datetime.now().isoformat()
    c.execute(
        'INSERT OR REPLACE INTO Cabot_Chalk '
        '(section_name, section_url, parent_section, depth, page_id, content_type, '
        'table_data_json, raw_text, links_json, sub_sections_json, tab_data_json, last_fetched) '
        'VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
        (
            section_name, section_url, parent_section, depth, page_id, content_type,
            json.dumps(table_data, ensure_ascii=False) if table_data else None,
            raw_text,
            json.dumps(links, ensure_ascii=False) if links else None,
            json.dumps(sub_sections, ensure_ascii=False) if sub_sections else None,
            json.dumps(tab_data, ensure_ascii=False) if tab_data else None,
            now
        )
    )
    c.commit()
    c.close()


def get_cabot_row_count():
    c = _conn()
    try:
        r = c.execute('SELECT COUNT(*) FROM Cabot_Chalk').fetchone()[0]
    except Exception:
        r = 0
    c.close()
    return r


def load_all_cabot_data():
    c = _conn()
    rows = c.execute('SELECT * FROM Cabot_Chalk ORDER BY depth, id').fetchall()
    c.close()
    return [dict(r) for r in rows]


# ================================================================
# SCRAPING HELPERS
# ================================================================

def scrape_tables(page):
    """Extract all tables from the main content area."""
    return page.evaluate(r"""() => {
        const tables = [];
        const contentArea = document.querySelector('#main-content') ||
                            document.querySelector('.wiki-content') ||
                            document.querySelector('[data-testid="page-content"]') ||
                            document.body;
        contentArea.querySelectorAll('table').forEach(tbl => {
            const headers = [];
            tbl.querySelectorAll('thead th, tr:first-child th').forEach(th => {
                headers.push(th.innerText.trim());
            });
            if (headers.length === 0) {
                const firstRow = tbl.querySelector('tr');
                if (firstRow) {
                    firstRow.querySelectorAll('td, th').forEach(cell => {
                        headers.push(cell.innerText.trim());
                    });
                }
            }
            const rows = [];
            const allRows = tbl.querySelectorAll('tr');
            const startIdx = headers.length > 0 ? 1 : 0;
            for (let i = startIdx; i < allRows.length; i++) {
                const cells = allRows[i].querySelectorAll('td, th');
                const row = {};
                cells.forEach((cell, ci) => {
                    const key = ci < headers.length ? headers[ci] : ('col_' + ci);
                    row[key] = cell.innerText.trim();
                });
                if (Object.values(row).some(v => v)) rows.push(row);
            }
            if (rows.length > 0) tables.push({headers: headers, rows: rows});
        });
        return tables;
    }""")


def scrape_raw_text(page):
    """Get all visible text from the main content area."""
    return page.evaluate(r"""() => {
        const el = document.querySelector('#main-content') ||
                   document.querySelector('.wiki-content') ||
                   document.querySelector('[data-testid="page-content"]') ||
                   document.body;
        return el ? el.innerText.trim() : '';
    }""")


def scrape_links(page):
    """Extract all links from the main content area."""
    return page.evaluate(r"""() => {
        const links = [];
        const el = document.querySelector('#main-content') ||
                   document.querySelector('.wiki-content') ||
                   document.querySelector('[data-testid="page-content"]') ||
                   document.body;
        el.querySelectorAll('a[href]').forEach(a => {
            const text = a.innerText.trim();
            const href = a.href;
            if (text && href && !href.startsWith('javascript:')) {
                links.push({text: text, href: href});
            }
        });
        return links;
    }""")


def expand_all_sections(page):
    """
    Expand ALL collapsed/expandable sections on the current page:
    - Confluence expand macros (click-to-expand sections)
    - Tab groups (click each tab to reveal content)
    - Accordion / toggle sections
    - Details/summary HTML elements
    Returns the number of elements expanded.
    """
    expanded = page.evaluate(r"""() => {
        let count = 0;

        // 1. Confluence expand macros
        document.querySelectorAll(
            '.expand-control:not(.expanded), ' +
            '.expand-control-text, ' +
            '.aui-expander-trigger[aria-expanded="false"], ' +
            'button.expand-control[aria-expanded="false"]'
        ).forEach(el => {
            try { el.click(); count++; } catch(e) {}
        });

        // 2. Details/summary elements
        document.querySelectorAll('details:not([open])').forEach(el => {
            try { el.setAttribute('open', ''); count++; } catch(e) {}
        });

        // 3. Confluence UI tabs — click each tab to load its content
        document.querySelectorAll(
            '.aui-tabs .menu-item a, ' +
            '.tabs-menu li a, ' +
            '[role="tab"][aria-selected="false"]'
        ).forEach(el => {
            try { el.click(); count++; } catch(e) {}
        });

        // 4. Toggle / accordion buttons
        document.querySelectorAll(
            'button[aria-expanded="false"], ' +
            '.toggle-trigger[aria-expanded="false"]'
        ).forEach(el => {
            // Skip sidebar tree toggles — only expand in-page content
            if (el.closest('.ia-splitter-left') || el.closest('nav') ||
                el.closest('.plugin_pagetree')) return;
            try { el.click(); count++; } catch(e) {}
        });

        return count;
    }""")
    return expanded


def scrape_tab_content(page):
    """
    If the page has tab groups (Confluence tabs macro), scrape each tab's content.
    Returns a list of {tab_name, tables, text, links} for each tab.
    """
    tab_data = page.evaluate(r"""() => {
        const results = [];
        const tabGroups = document.querySelectorAll(
            '.aui-tabs, .tabs-pane-container, [role="tablist"]'
        );

        // Also check for Confluence-style tabs
        const confTabs = document.querySelectorAll('.tabs-menu');

        const allTabContainers = new Set();
        tabGroups.forEach(tg => allTabContainers.add(tg.closest('.aui-tabs') || tg));
        confTabs.forEach(ct => allTabContainers.add(ct.closest('.aui-tabs') || ct.parentElement));

        allTabContainers.forEach(container => {
            if (!container) return;
            const tabs = container.querySelectorAll(
                '.tabs-menu li a, .menu-item a, [role="tab"]'
            );
            const panes = container.querySelectorAll(
                '.tabs-pane, [role="tabpanel"]'
            );

            tabs.forEach((tab, idx) => {
                const tabName = tab.innerText.trim();
                if (!tabName) return;

                // Click the tab to make its pane visible
                try { tab.click(); } catch(e) {}

                // Find the corresponding pane
                let pane = panes[idx] || null;
                if (!pane) {
                    // Try to find by aria-controls
                    const controls = tab.getAttribute('aria-controls');
                    if (controls) pane = document.getElementById(controls);
                }
                if (!pane) return;

                // Scrape tables from this pane
                const tables = [];
                pane.querySelectorAll('table').forEach(tbl => {
                    const headers = [];
                    tbl.querySelectorAll('thead th, tr:first-child th').forEach(th => {
                        headers.push(th.innerText.trim());
                    });
                    if (headers.length === 0) {
                        const firstRow = tbl.querySelector('tr');
                        if (firstRow) {
                            firstRow.querySelectorAll('td, th').forEach(cell => {
                                headers.push(cell.innerText.trim());
                            });
                        }
                    }
                    const rows = [];
                    const allRows = tbl.querySelectorAll('tr');
                    const startIdx = headers.length > 0 ? 1 : 0;
                    for (let i = startIdx; i < allRows.length; i++) {
                        const cells = allRows[i].querySelectorAll('td, th');
                        const row = {};
                        cells.forEach((cell, ci) => {
                            const key = ci < headers.length ? headers[ci] : ('col_' + ci);
                            row[key] = cell.innerText.trim();
                        });
                        if (Object.values(row).some(v => v)) rows.push(row);
                    }
                    if (rows.length > 0) tables.push({headers, rows});
                });

                // Scrape text
                const text = pane.innerText ? pane.innerText.trim() : '';

                // Scrape links
                const links = [];
                pane.querySelectorAll('a[href]').forEach(a => {
                    const lt = a.innerText.trim();
                    const lh = a.href;
                    if (lt && lh && !lh.startsWith('javascript:')) {
                        links.push({text: lt, href: lh});
                    }
                });

                results.push({
                    tab_name: tabName,
                    tables: tables,
                    text: text.substring(0, 50000),
                    links: links
                });
            });
        });

        return results;
    }""")
    return tab_data


def extract_page_id_from_url(url):
    """Extract the numeric page ID from a Confluence URL."""
    m = re.search(r'/pages/(\d+)', url or '')
    return m.group(1) if m else None


def expand_cabot_tree(page):
    """Expand the Cabot node and ALL its children in the sidebar tree."""
    for attempt in range(8):
        expanded = page.evaluate(r"""() => {
            let c = 0;
            const PAGE_ID = '3247714982';

            // Find the Cabot anchor by page ID
            let anchor = null;
            document.querySelectorAll('a[href]').forEach(a => {
                if (a.href && a.href.indexOf(PAGE_ID) !== -1) anchor = a;
            });
            if (!anchor) return 0;

            // Walk up to the owning <li>
            let parentLi = anchor.closest('li');
            if (!parentLi) return 0;

            // Expand ALL toggle buttons within this entire subtree (deep)
            parentLi.querySelectorAll(
                'button[aria-expanded="false"], ' +
                'a.plugin_pagetree_childtoggle_haschildren'
            ).forEach(btn => {
                try { btn.click(); c++; } catch(e) {}
            });

            return c;
        }""")
        if expanded:
            log('    Expanded %d Cabot tree nodes (pass %d)' % (expanded, attempt + 1))
            time.sleep(2)
        else:
            break
    time.sleep(1)


def discover_all_children(page, parent_page_id, space_key=SPACE_KEY):
    """
    Find ALL child pages under a given parent page ID in the sidebar tree.
    Works recursively — returns children at all levels under the parent.
    """
    items = page.evaluate(r"""(args) => {
        const items = [];
        const seen = new Set();
        const parentId = args.parentId;
        const spaceKey = args.spaceKey;
        const spacePattern = '/spaces/' + spaceKey + '/pages/';

        // Strategy 1: Anchor to parent <li> in page-tree
        let anchor = null;
        document.querySelectorAll('a[href]').forEach(a => {
            if (a.href && a.href.indexOf(parentId) !== -1) anchor = a;
        });

        if (anchor) {
            let parentLi = anchor.closest('li');
            if (parentLi) {
                // Get ALL descendant links (not just direct children)
                parentLi.querySelectorAll('a[href*="' + spacePattern + '"]').forEach(a => {
                    const text = a.innerText.trim();
                    const href = a.href;
                    if (text && href && !seen.has(href) &&
                        href.indexOf(parentId) === -1) {
                        seen.add(href);
                        items.push({text: text, href: href});
                    }
                });
            }
        }

        // Strategy 2: Confluence Server plugin_pagetree
        if (items.length === 0) {
            let currentNode = document.querySelector('.plugin_pagetree_current') ||
                              document.querySelector('[data-pageid="' + parentId + '"]');
            if (currentNode) {
                let parentLi = currentNode.closest('li') || currentNode.parentElement;
                let childContainer = parentLi ?
                    parentLi.querySelector('.plugin_pagetree_children_container') : null;
                if (childContainer) {
                    childContainer.querySelectorAll('a[href*="' + spacePattern + '"]').forEach(a => {
                        const text = a.innerText.trim();
                        const href = a.href;
                        if (text && href && !seen.has(href) && href.indexOf(parentId) === -1) {
                            seen.add(href);
                            items.push({text: text, href: href});
                        }
                    });
                }
            }
        }

        // Strategy 3: Depth-based walk
        if (items.length === 0) {
            const allLinks = document.querySelectorAll(
                '.ia-splitter-left a[href*="' + spacePattern + '"], ' +
                'nav a[href*="' + spacePattern + '"]'
            );
            let found = false;
            let parentDepth = -1;

            function liDepth(el) {
                let d = 0;
                while (el) { if (el.tagName === 'LI') d++; el = el.parentElement; }
                return d;
            }

            allLinks.forEach(a => {
                const text = a.innerText.trim();
                const href = a.href;
                if (!found) {
                    if (href.indexOf(parentId) !== -1) {
                        found = true;
                        parentDepth = liDepth(a);
                    }
                    return;
                }
                const d = liDepth(a);
                if (d > parentDepth && text && href && !seen.has(href)) {
                    seen.add(href);
                    items.push({text: text, href: href});
                } else if (d <= parentDepth) {
                    found = false;
                }
            });
        }

        return items;
    }""", {'parentId': parent_page_id, 'spaceKey': space_key})

    return items


def discover_in_page_children(page, space_key=SPACE_KEY):
    """
    Discover child page links embedded IN the page content itself:
    - Confluence children macro
    - Bullet-list links to sub-pages
    - Any links pointing to pages in the same space
    """
    items = page.evaluate(r"""(spaceKey) => {
        const items = [];
        const seen = new Set();
        const spacePattern = '/spaces/' + spaceKey + '/pages/';
        const contentArea = document.querySelector('#main-content') ||
                            document.querySelector('.wiki-content') ||
                            document.querySelector('[data-testid="page-content"]') ||
                            document.body;

        // Children macro links
        contentArea.querySelectorAll(
            '.childpages-macro a, .children-show-if a, ' +
            '.plugin_pagetree_children_container a'
        ).forEach(a => {
            const text = a.innerText.trim();
            const href = a.href;
            if (text && href && !seen.has(href) && href.indexOf(spacePattern) !== -1) {
                seen.add(href);
                items.push({text: text, href: href});
            }
        });

        // Also pick up any bullet-list links to same-space pages
        contentArea.querySelectorAll('li a[href*="' + spacePattern + '"]').forEach(a => {
            const text = a.innerText.trim();
            const href = a.href;
            if (text && href && !seen.has(href)) {
                seen.add(href);
                items.push({text: text, href: href});
            }
        });

        return items;
    }""", space_key)
    return items


# ================================================================
# DOCX GENERATION
# ================================================================

def generate_docx(all_data):
    """Generate a DOCX document with all scraped Cabot data."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        log('[DOCX] python-docx not installed. Run: pip install python-docx')
        log('[DOCX] Skipping DOCX generation.')
        return None

    doc = Document()
    doc.add_heading('Cabot Chalk — API Documentation Export', level=0)
    doc.add_paragraph('Generated: %s' % datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    doc.add_paragraph('Source: %s' % CABOT_URL)
    doc.add_paragraph('Total sections scraped: %d' % len(all_data))
    doc.add_paragraph('')

    for section in all_data:
        name = section.get('section_name', 'Unknown')
        parent = section.get('parent_section', '')
        depth = section.get('depth', 0)
        url = section.get('section_url', '')
        raw = section.get('raw_text', '')
        table_json = section.get('table_data_json')
        links_json = section.get('links_json')
        subs_json = section.get('sub_sections_json')
        tab_json = section.get('tab_data_json')

        # Heading level based on depth (1-4)
        hlevel = min(depth + 1, 4)
        indent = '  ' * depth
        heading_text = '%s%s' % (indent, name) if depth > 0 else name
        if parent and parent != name and parent != 'Cabot':
            heading_text = '%s > %s' % (parent, name)
        doc.add_heading(heading_text, level=hlevel)

        if url:
            doc.add_paragraph('URL: %s' % url)
        if depth > 0:
            doc.add_paragraph('Depth: %d | Parent: %s' % (depth, parent))

        # Tables
        tables = []
        if table_json:
            try:
                tables = json.loads(table_json) if isinstance(table_json, str) else table_json
            except Exception:
                pass

        if tables:
            for ti, tbl in enumerate(tables):
                if ti > 0:
                    doc.add_paragraph('')
                headers = tbl.get('headers', [])
                rows = tbl.get('rows', [])
                if not rows:
                    continue
                cols = headers if headers else list(rows[0].keys()) if rows else []
                if not cols:
                    continue

                doc.add_paragraph('Table %d (%d rows):' % (ti + 1, len(rows)))
                tbl_obj = doc.add_table(rows=1 + len(rows), cols=len(cols))
                tbl_obj.style = 'Table Grid'

                for ci, col in enumerate(cols):
                    cell = tbl_obj.rows[0].cells[ci]
                    cell.text = str(col)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

                for ri, row in enumerate(rows):
                    for ci, col in enumerate(cols):
                        tbl_obj.rows[ri + 1].cells[ci].text = str(row.get(col, ''))

        # Tab content
        tab_data = []
        if tab_json:
            try:
                tab_data = json.loads(tab_json) if isinstance(tab_json, str) else tab_json
            except Exception:
                pass

        if tab_data:
            doc.add_heading('Tab Content', level=min(hlevel + 1, 6))
            for tab in tab_data:
                tab_name = tab.get('tab_name', 'Tab')
                doc.add_heading('Tab: %s' % tab_name, level=min(hlevel + 2, 6))

                tab_tables = tab.get('tables', [])
                for tti, ttbl in enumerate(tab_tables):
                    headers = ttbl.get('headers', [])
                    rows = ttbl.get('rows', [])
                    if not rows:
                        continue
                    cols = headers if headers else list(rows[0].keys()) if rows else []
                    if not cols:
                        continue
                    doc.add_paragraph('Table %d (%d rows):' % (tti + 1, len(rows)))
                    tbl_obj = doc.add_table(rows=1 + len(rows), cols=len(cols))
                    tbl_obj.style = 'Table Grid'
                    for ci, col in enumerate(cols):
                        cell = tbl_obj.rows[0].cells[ci]
                        cell.text = str(col)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
                    for ri, row in enumerate(rows):
                        for ci, col in enumerate(cols):
                            tbl_obj.rows[ri + 1].cells[ci].text = str(row.get(col, ''))

                tab_text = tab.get('text', '')
                if tab_text and not tab_tables:
                    doc.add_paragraph(tab_text[:3000])

        # Links
        links = []
        if links_json:
            try:
                links = json.loads(links_json) if isinstance(links_json, str) else links_json
            except Exception:
                pass

        if links:
            doc.add_heading('Links', level=min(hlevel + 1, 6))
            for lnk in links:
                doc.add_paragraph('\u2022 %s \u2014 %s' % (lnk.get('text', ''), lnk.get('href', '')))

        # Sub-sections
        subs = []
        if subs_json:
            try:
                subs = json.loads(subs_json) if isinstance(subs_json, str) else subs_json
            except Exception:
                pass

        if subs:
            doc.add_heading('Sub-sections', level=min(hlevel + 1, 6))
            for sub in subs:
                doc.add_paragraph('\u2022 %s \u2014 %s' % (sub.get('text', ''), sub.get('href', '')))

        # Raw text fallback (when no tables)
        if raw and not tables and not tab_data:
            doc.add_heading('Content', level=min(hlevel + 1, 6))
            text_to_add = raw[:5000]
            if len(raw) > 5000:
                text_to_add += '\n... [truncated \u2014 see DB for full text]'
            doc.add_paragraph(text_to_add)

        doc.add_page_break()

    DOCX_PATH.parent.mkdir(exist_ok=True)
    doc.save(str(DOCX_PATH))
    log('  [DOCX] Saved: %s' % DOCX_PATH)
    return DOCX_PATH


# ================================================================
# RECURSIVE PAGE SCRAPER
# ================================================================

def scrape_page(page, page_url, page_name, parent_name, depth, max_depth,
                visited, stats):
    """
    Scrape a single Chalk page and recursively scrape all its sub-pages.

    Args:
        page:        Playwright page object
        page_url:    URL to navigate to
        page_name:   Display name of this page
        parent_name: Name of the parent section
        depth:       Current recursion depth
        max_depth:   Maximum recursion depth
        visited:     Set of already-visited URLs (prevents cycles)
        stats:       Dict with 'scraped' and 'failed' counters
    """
    # Normalize URL for dedup
    norm_url = page_url.split('?')[0].split('#')[0].rstrip('/')
    if norm_url in visited:
        log('%s[SKIP] Already visited: %s' % ('  ' * depth, page_name))
        return
    visited.add(norm_url)

    if depth > max_depth:
        log('%s[SKIP] Max depth %d reached for: %s' % ('  ' * depth, max_depth, page_name))
        return

    indent = '  ' * depth
    page_id = extract_page_id_from_url(page_url)

    log('%s[D%d] Scraping: %s' % (indent, depth, page_name))
    log('%s  URL: %s' % (indent, page_url[:120]))

    try:
        page.goto(page_url, wait_until='domcontentloaded', timeout=PAGE_LOAD_TIMEOUT_MS)
        _wait(page)
        time.sleep(3)

        actual_url = page.url
        log('%s  Loaded: %s' % (indent, actual_url[:120]))

        # Expand all collapsed sections, tabs, accordions on this page
        exp_count = expand_all_sections(page)
        if exp_count:
            log('%s  Expanded %d sections/tabs' % (indent, exp_count))
            time.sleep(1)
            # Second pass — some sections reveal more expandables
            exp2 = expand_all_sections(page)
            if exp2:
                log('%s  Expanded %d more sections (pass 2)' % (indent, exp2))
                time.sleep(1)

        # Scrape main content
        tables = scrape_tables(page)
        raw_text = scrape_raw_text(page)
        links = scrape_links(page)

        # Scrape tab content if any
        tab_data = scrape_tab_content(page)

        # Discover sub-pages from in-page content (children macro, bullet links)
        in_page_children = discover_in_page_children(page)

        content_type = 'table' if tables else ('tabs' if tab_data else 'text')
        table_count = len(tables)
        row_count = sum(len(t.get('rows', [])) for t in tables)
        tab_count = len(tab_data)

        save_cabot_section(
            section_name=page_name,
            section_url=page_url,
            parent_section=parent_name,
            depth=depth,
            page_id=page_id,
            content_type=content_type,
            table_data=tables,
            raw_text=raw_text[:50000],
            links=links,
            sub_sections=in_page_children if in_page_children else None,
            tab_data=tab_data if tab_data else None
        )

        stats['scraped'] += 1
        log('%s  OK — %d table(s), %d rows, %d tabs, %d links, %d sub-page refs' % (
            indent, table_count, row_count, tab_count, len(links), len(in_page_children)))

        # Recursively scrape sub-pages found in the content
        if in_page_children and depth < max_depth:
            log('%s  Diving into %d sub-pages...' % (indent, len(in_page_children)))
            for child in in_page_children:
                child_name = child.get('text', '').strip()
                child_url = child.get('href', '')
                if not child_name or not child_url:
                    continue
                # Only follow links within the same space
                if '/spaces/%s/pages/' % SPACE_KEY not in child_url:
                    log('%s  [SKIP] Not in %s space: %s' % (indent, SPACE_KEY, child_name))
                    continue
                scrape_page(page, child_url, child_name, page_name,
                            depth + 1, max_depth, visited, stats)

    except Exception as e:
        stats['failed'] += 1
        log('%s  FAILED: %s' % (indent, str(e)[:150]))


# ================================================================
# MAIN SCRAPER
# ================================================================

def scrape_cabot(force=False, headless=True, max_depth=DEFAULT_MAX_DEPTH):
    init_cabot_table()

    if not force and get_cabot_row_count() > 0:
        count = get_cabot_row_count()
        log('DB already has %d Cabot sections. Use --force to re-fetch.' % count)
        log('Generating DOCX from existing data...')
        all_data = load_all_cabot_data()
        generate_docx(all_data)
        return

    # If forcing, clear existing data
    if force:
        c = _conn()
        c.execute('DELETE FROM Cabot_Chalk')
        c.commit()
        c.close()
        log('  [DB] Cleared existing Cabot_Chalk data for fresh fetch.')

    t0 = time.time()
    log('=' * 70)
    log('  CABOT CHALK DEEP SCRAPER')
    log('  Target:    %s' % CABOT_URL)
    log('  Space:     %s' % SPACE_KEY)
    log('  Page ID:   %s' % CABOT_PAGE_ID)
    log('  DB:        %s' % DB_PATH)
    log('  Max Depth: %d' % max_depth)
    log('=' * 70)

    log('\nLaunching browser...')
    pw = sync_playwright().start()
    br = pw.chromium.launch(headless=headless, channel=get_browser_channel())
    ctx = br.new_context(viewport={'width': 1920, 'height': 1080})
    page = ctx.new_page()

    visited = set()
    stats = {'scraped': 0, 'failed': 0}

    # ---- Step 1: Load the main Cabot page ----
    log('\n[1/4] Loading Cabot main page...')
    page.goto(CABOT_URL, wait_until='domcontentloaded', timeout=PAGE_LOAD_TIMEOUT_MS)
    _wait(page)
    time.sleep(5)
    log('  Page title: %s' % page.title()[:80])

    current_url = page.url
    if CABOT_PAGE_ID not in current_url:
        log('  WARNING: Current URL does not contain Cabot page ID!')
        log('  Current URL: %s' % current_url)
        log('  Expected page ID: %s' % CABOT_PAGE_ID)

    # ---- Step 2: Scrape the main Cabot page content ----
    log('\n[2/4] Scraping main Cabot page content...')
    exp = expand_all_sections(page)
    if exp:
        log('  Expanded %d sections on main page' % exp)
        time.sleep(1)

    main_tables = scrape_tables(page)
    main_text = scrape_raw_text(page)
    main_links = scrape_links(page)
    main_tabs = scrape_tab_content(page)
    main_in_page = discover_in_page_children(page)

    log('  Found %d table(s), %d tabs, %d links, %d in-page child refs' % (
        len(main_tables), len(main_tabs), len(main_links), len(main_in_page)))

    save_cabot_section(
        section_name='Cabot (Main)',
        section_url=CABOT_URL,
        parent_section='Cabot',
        depth=0,
        page_id=CABOT_PAGE_ID,
        content_type='table' if main_tables else 'text',
        table_data=main_tables,
        raw_text=main_text[:50000],
        links=main_links,
        sub_sections=main_in_page if main_in_page else None,
        tab_data=main_tabs if main_tabs else None
    )
    stats['scraped'] += 1
    visited.add(CABOT_URL.split('?')[0].split('#')[0].rstrip('/'))
    log('  Saved main page data.')

    # ---- Step 3: Expand sidebar tree and discover ALL children ----
    log('\n[3/4] Expanding Cabot sidebar tree (deep)...')
    expand_cabot_tree(page)
    sidebar_items = discover_all_children(page, CABOT_PAGE_ID)
    log('  Found %d items in sidebar tree under Cabot' % len(sidebar_items))

    for item in sidebar_items:
        log('    - %s  (%s)' % (item.get('text', ''), item.get('href', '')[:80]))

    if not sidebar_items and not main_in_page:
        log('  WARNING: No child items found! The page tree may not have loaded.')
        log('  Try running with --headless false to debug.')

    # Merge sidebar items with in-page children (dedup by URL)
    all_children = {}
    for item in sidebar_items:
        url = item.get('href', '').split('?')[0].split('#')[0].rstrip('/')
        if url:
            all_children[url] = item
    for item in main_in_page:
        url = item.get('href', '').split('?')[0].split('#')[0].rstrip('/')
        if url and url not in all_children:
            all_children[url] = item

    merged_children = list(all_children.values())
    log('  Merged total: %d unique child pages to scrape' % len(merged_children))

    # ---- Step 4: Deep-scrape each child page recursively ----
    log('\n[4/4] Deep-scraping all child pages (max depth=%d)...' % max_depth)

    for idx, item in enumerate(merged_children, 1):
        item_name = item.get('text', '').strip()
        item_url = item.get('href', '')

        if not item_name or not item_url:
            continue

        # Safety: only follow links within the SPECTRUMMOBIT2 space
        if '/spaces/%s/pages/' % SPACE_KEY not in item_url:
            log('  [%d/%d] SKIP (not in %s space): %s' % (
                idx, len(merged_children), SPACE_KEY, item_name))
            continue

        log('\n  === [%d/%d] %s ===' % (idx, len(merged_children), item_name))
        scrape_page(page, item_url, item_name, 'Cabot',
                    depth=1, max_depth=max_depth,
                    visited=visited, stats=stats)

    # Cleanup browser
    ctx.close()
    br.close()
    pw.stop()

    # ---- Generate DOCX ----
    log('\nGenerating DOCX report...')
    all_data = load_all_cabot_data()
    generate_docx(all_data)

    # Summary
    elapsed = time.time() - t0
    m, s = divmod(int(elapsed), 60)
    log('\n' + '=' * 70)
    log('  CABOT CHALK DEEP SCRAPE COMPLETE')
    log('  Sections scraped: %d' % stats['scraped'])
    log('  Failed:           %d' % stats['failed'])
    log('  DB rows:          %d' % get_cabot_row_count())
    log('  DB path:          %s' % DB_PATH)
    log('  DOCX path:        %s' % DOCX_PATH)
    log('  Time:             %dm %ds' % (m, s))
    log('=' * 70)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Deep-scrape Cabot Chalk pages into CABOT_CHALK_DB + DOCX')
    parser.add_argument('--force', action='store_true',
                        help='Re-fetch even if DB has data')
    parser.add_argument('--headless', type=str, default='true',
                        help='Run browser headless (true/false). Default: true')
    parser.add_argument('--max-depth', type=int, default=DEFAULT_MAX_DEPTH,
                        help='Maximum recursion depth for sub-pages. Default: %d' % DEFAULT_MAX_DEPTH)
    args = parser.parse_args()
    headless = args.headless.lower() != 'false'
    scrape_cabot(force=args.force, headless=headless, max_depth=args.max_depth)
