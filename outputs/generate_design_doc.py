"""
Generate a polished DOCX Design & Implementation document for TSG, TSE, MDA dashboards.
Run: python TestSuiteGenerator/outputs/generate_design_doc.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import datetime

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.font.name = 'Calibri'

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Automation Dashboards')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('TSG  |  TSE  |  MDA')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph('')

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Design & Implementation Guide')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
doc.add_paragraph('')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Version 1.0  •  April 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Our QA team spends a significant portion of each sprint on repetitive tasks: '
    'writing test cases from Jira stories, manually running API tests, capturing screenshots '
    'for evidence, and pulling weekly status reports from QMetry and Jira. '
    'These three automation dashboards eliminate that manual work and create an end-to-end '
    'pipeline: Generate → Execute → Report.'
)

add_table(doc,
    ['Dashboard', 'What It Does', 'Time Saved'],
    [
        ['TSG (Test Suite Generator)', 'Reads Jira features and writes test cases automatically', 'Hours per feature → minutes'],
        ['TSE (Test Suite Executor)', 'Runs test cases against live APIs and captures proof', 'Full day of manual testing → automated'],
        ['MDA (Jira Dashboard)', 'Pulls QMetry/Jira data and builds the weekly status deck', '2-3 hours of copy-paste → one click'],
    ]
)

doc.add_paragraph('')

# ============================================================
# 2. WHY THIS MATTERS
# ============================================================
doc.add_heading('2. Why This Matters', level=1)

doc.add_heading('Before Automation', level=2)
bullets_before = [
    'Test engineers manually read Jira stories, Chalk docs, and attachments to write test cases in Excel — often missing edge cases or duplicating effort across team members.',
    'Executing tests meant stepping through each API call by hand, taking screenshots, and assembling evidence documents.',
    'Weekly reporting required logging into QMetry, exporting data, filtering in Excel, building pivot tables, and pasting into PowerPoint.',
]
for b in bullets_before:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('After Automation', level=2)
bullets_after = [
    'TSG generates 100+ test cases per feature in minutes, covering happy paths, error scenarios, and integration checks.',
    'TSE executes every test case automatically with full evidence capture — screenshots, API responses, and validation results packaged into a Word document.',
    'MDA produces the weekly status deck in one click — QMetry exports, Jira defect summaries, and execution reports all flow into a ready-to-present PowerPoint.',
]
for b in bullets_after:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Key Benefits', level=2)
benefits = [
    'Consistent test coverage across all features (no missed scenarios)',
    'Complete audit trail for every test execution',
    'Faster sprint cycles — less time on paperwork, more time on actual testing',
    'Vendor-aware reporting for multi-team coordination',
    'Standardized output formats (QMetry Excel, Evidence DOCX, Status PPTX)',
]
for b in benefits:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# 3. HOW EACH DASHBOARD WORKS
# ============================================================
doc.add_heading('3. How Each Dashboard Works', level=1)

# --- TSG ---
doc.add_heading('3.1  TSG — Test Suite Generator', level=2)
doc.add_paragraph(
    'Purpose: Automatically create comprehensive test suites from Jira feature specifications.'
)

doc.add_heading('How It Works', level=3)
tsg_steps = [
    'You enter a Jira feature ID (e.g., MWTGPROV-3949).',
    'TSG pulls the feature details from Jira — acceptance criteria, description, linked issues.',
    'It scrapes the Chalk documentation page for business scenarios and rules.',
    'It checks any attached test data files for additional coverage.',
    'The test engine reasons through the feature like a senior QA engineer — identifying what needs positive testing, negative testing, boundary checks, and integration validation.',
    'It outputs a QMetry-compatible Excel file with structured test cases and steps, plus a Feature Summary document.',
]
for i, s in enumerate(tsg_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Key Capabilities', level=3)
tsg_caps = [
    'Pulls data from multiple sources (Jira, Chalk, attachments, linked issues)',
    'Classifies feature type (API, UI portal, notification, batch process) and adjusts test strategy',
    'Removes duplicate scenarios and filters out irrelevant data',
    'Tracks every generation with version history for audit',
    'Caches all data locally so the dashboard loads instantly after first run',
    'Optional AI-powered analysis for enhanced scenario reasoning',
]
for c in tsg_caps:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Core Modules', level=3)
add_table(doc,
    ['Module', 'Purpose'],
    [
        ['Test Engine', 'Builds the test suite structure, applies generation strategies'],
        ['Test Analyst', 'Reasons about feature type, generates lifecycle-specific scenarios'],
        ['Database', 'SQLite cache for features, Jira data, Chalk content, generation history'],
        ['Diff Engine', 'Compares test suite versions to show what changed'],
        ['QMetry Exporter', 'Produces formatted Excel files ready for QMetry import'],
        ['Linked Fetcher', 'Traverses Jira issue links to pull related acceptance criteria'],
    ]
)
doc.add_paragraph('')

# --- TSE ---
doc.add_heading('3.2  TSE — Test Suite Executor', level=2)
doc.add_paragraph(
    'Purpose: Automatically run generated test suites against live APIs and capture evidence.'
)

doc.add_heading('How It Works', level=3)
tse_steps = [
    'Load a test suite from the TSG database or an Excel file.',
    'TSE classifies each test case by type — activation flow, device change, UI check, etc.',
    'For each test step, TSE maps it to the right action: call an API, check a portal screen, validate a report.',
    'It executes a 7-step pipeline for each test case (see below).',
    'Results are packaged into a Word document with inline screenshots and pass/fail status.',
]
for i, s in enumerate(tse_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Execution Pipeline (7 Steps)', level=3)
pipeline = [
    ('OAuth', 'Obtain authentication token'),
    ('Device Validation', 'Verify test device is certified'),
    ('API Call', 'Execute the actual operation (activate, change, deactivate, etc.)'),
    ('Century Report', 'Pull transaction report for verification'),
    ('Service Grouping', 'Validate expected transactions against templates'),
    ('NBOP Check', 'Verify portal reflects the changes'),
    ('Evidence', 'Package screenshots and responses into the final document'),
]
add_table(doc, ['Step', 'Action'], [[f'{i+1}. {p[0]}', p[1]] for i, p in enumerate(pipeline)])
doc.add_paragraph('')

doc.add_heading('Key Capabilities', level=3)
tse_caps = [
    'Loads test suites from database or Excel (flexible input)',
    'Maps test step descriptions to executable API actions automatically',
    'Falls back to raw curl file execution for APIs without dedicated modules (100% coverage)',
    'Auto-allocates test devices from the SharePoint device pool',
    'Self-healing: refreshes expired tokens, retries failed steps, falls back to alternative paths',
    'Generates complete evidence documents for compliance and audit',
    'Intent-based execution (V2) adapts to test type automatically',
]
for c in tse_caps:
    doc.add_paragraph(c, style='List Bullet')

# --- MDA ---
doc.add_heading('3.3  MDA — Jira Dashboard', level=2)
doc.add_paragraph(
    'Purpose: Automate weekly QMetry/Jira data extraction and status report generation.'
)

doc.add_heading('How It Works', level=3)
mda_modules = [
    ('QMetry Export', 'Navigates to QMetry in Jira, selects the test folder, exports test cases to Excel, creates a pivot table grouped by labels and assignee.'),
    ('Jira Defect Extract', 'Runs a JQL search for defects, filters by reporter, creates a status/priority pivot.'),
    ('Test Execution Report', 'Fetches QMetry execution summary data and generates a styled Excel report.'),
    ('PowerPoint Generation', 'Auto-builds a weekly status deck with title slide, delivery updates, QMetry summary, defect summary, and execution overview.'),
]
for name, desc in mda_modules:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Key Capabilities', level=3)
mda_caps = [
    'Vendor-aware filtering (All Vendors, Infy team, Other Vendors) for multi-team reporting',
    'Handles QMetry UI quirks — lazy-loaded tree nodes, cold-session startup, virtual scrolling',
    'Produces presentation-ready PowerPoint with styled tables',
    'Single browser session for performance (no repeated logins)',
]
for c in mda_caps:
    doc.add_paragraph(c, style='List Bullet')

# ============================================================
# 4. HOW THEY CONNECT
# ============================================================
doc.add_heading('4. How They Connect', level=1)

doc.add_paragraph(
    'The three dashboards form a pipeline where each tool feeds into the next:'
)

add_table(doc,
    ['From', 'To', 'What Flows'],
    [
        ['TSG', 'TSE', 'Test suites (via SQLite database or Excel files)'],
        ['TSE', 'MDA', 'Execution results feed into the weekly reporting cycle'],
        ['All', 'All', 'Shared technology stack: Python, Streamlit, Playwright, SQLite'],
    ]
)
doc.add_paragraph('')

# ============================================================
# 5. TECHNOLOGY STACK
# ============================================================
doc.add_heading('5. Technology Stack', level=1)

add_table(doc,
    ['Component', 'Technology', 'Why'],
    [
        ['Language', 'Python 3.9+', 'Standard, well-supported, rich library ecosystem'],
        ['Dashboard UI', 'Streamlit', 'Fast to build, interactive, no frontend code needed'],
        ['Browser Automation', 'Playwright', 'Reliable, cross-browser, handles modern web apps'],
        ['Data Storage', 'SQLite (WAL mode)', 'Zero setup, ships with Python, concurrent reads'],
        ['Excel Processing', 'openpyxl', 'Read/write Excel files with formatting'],
        ['PowerPoint', 'python-pptx', 'Programmatic slide generation'],
        ['Data Analysis', 'pandas', 'Pivot tables, filtering, data transformation'],
        ['HTTP Client', 'requests', 'API calls to Jira, QMetry, TMO APIs'],
        ['AI (optional)', 'OpenAI / Azure / Bedrock', 'Enhanced test case reasoning in TSG'],
    ]
)
doc.add_paragraph('')

# ============================================================
# 6. EXTERNAL SYSTEMS
# ============================================================
doc.add_heading('6. External Systems', level=1)

add_table(doc,
    ['System', 'Used By', 'Purpose'],
    [
        ['Jira', 'TSG, MDA', 'Feature specs, acceptance criteria, defect tracking'],
        ['Chalk', 'TSG', 'Business scenario documentation'],
        ['QMetry', 'TSG, MDA', 'Test case management, execution tracking'],
        ['TMO APIs', 'TSE', 'API execution (activate, change, deactivate, etc.)'],
        ['NBOP Portal', 'TSE', 'UI validation and field verification'],
        ['Century Report', 'TSE', 'Transaction verification'],
        ['SharePoint', 'TSG, TSE', 'Test device and SIM data'],
        ['ALM Octane', 'OTM (optional)', 'Test management for Octane-based projects'],
    ]
)
doc.add_paragraph('')

# ============================================================
# 7. SETUP & INSTALLATION
# ============================================================
doc.add_heading('7. Setup & Installation', level=1)

doc.add_heading('Prerequisites', level=2)
prereqs = [
    'Python 3.9 or higher',
    'Jira account with API access',
    'Network access to Chalk, QMetry, and TMO environments',
    'Chrome, Edge, or Chromium browser installed',
]
for p in prereqs:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('Step 1: Install Dependencies', level=2)
p = doc.add_paragraph()
run = p.add_run('pip install streamlit playwright pandas openpyxl python-pptx requests lxml\nplaywright install chromium')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Step 2: Configure Credentials', level=2)
doc.add_paragraph('Copy .env.example to .env and fill in your Jira credentials:')
p = doc.add_paragraph()
run = p.add_run('JIRA_USER=your_jira_username\nJIRA_PASS=your_jira_password\nHEADLESS=false')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('Step 3: Pre-load Data (TSG, one-time)', level=2)
p = doc.add_paragraph()
run = p.add_run('python TestSuiteGenerator/preload_db.py')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('This fetches all PI features and Chalk data into the local database. Takes a few minutes on first run, then the dashboard loads instantly.')

doc.add_heading('Step 4: Launch a Dashboard', level=2)
cmds = [
    ('Test Suite Generator', 'streamlit run TSG_Dashboard_V4.1.py'),
    ('Test Suite Executor', 'streamlit run TestSuiteExecutor/TSE_Dashboard_V1.0.py'),
    ('MDA Jira Dashboard', 'streamlit run MDA_Jira_Dashboard_V5.1.py'),
]
for name, cmd in cmds:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    run = p.add_run(cmd)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

# ============================================================
# 8. DESIGN DECISIONS
# ============================================================
doc.add_heading('8. Design Decisions', level=1)

add_table(doc,
    ['Decision', 'Rationale'],
    [
        ['SQLite over a database server', 'Zero setup, single file, ships with Python. WAL mode handles concurrent reads. No DBA needed.'],
        ['Streamlit for the UI', 'Rapid development, interactive widgets, no HTML/CSS/JS required.'],
        ['Playwright over Selenium', 'More reliable, better SPA handling, built-in wait strategies, cross-browser support.'],
        ['Local caching everywhere', 'Jira and Chalk are slow to query. Local cache means dashboards load in seconds. Stale data warnings after 24 hours.'],
        ['Curl fallback in TSE', 'Not every API has a dedicated module. Parsing raw curl files ensures 100% test coverage.'],
        ['Modular architecture', 'Each dashboard is split into focused modules. Easy to maintain, test, and extend.'],
        ['Intent-based execution (V2)', 'Different test types need different strategies. Classifying by intent lets the engine pick the right handler.'],
    ]
)
doc.add_paragraph('')

# ============================================================
# 9. SECURITY
# ============================================================
doc.add_heading('9. Security Considerations', level=1)
security = [
    'Credentials stored in .env files, never committed to source control.',
    'Jira authentication uses existing corporate credentials.',
    'OAuth tokens refreshed automatically and never persisted to disk.',
    'SQLite databases are local-only — no network exposure.',
    'Browser automation runs with CREATE_NO_WINDOW on Windows to suppress popups.',
]
for s in security:
    doc.add_paragraph(s, style='List Bullet')

# ============================================================
# 10. SUMMARY
# ============================================================
doc.add_heading('10. Impact Summary', level=1)

doc.add_paragraph(
    'These three dashboards transform the QA workflow from manual, error-prone processes '
    'into an automated pipeline. TSG eliminates test case writing. TSE eliminates manual '
    'test execution. MDA eliminates report assembly. Together, they free the team to focus '
    'on what matters — finding real issues and improving quality.'
)

add_table(doc,
    ['Metric', 'Before', 'After'],
    [
        ['Test case creation', '2-4 hours per feature', '5 minutes'],
        ['Test execution + evidence', 'Full day', 'Automated'],
        ['Weekly status report', '2-3 hours', 'One click'],
        ['Test coverage consistency', 'Varies by person', 'Standardized'],
        ['Evidence documentation', 'Manual screenshots', 'Auto-generated Word doc'],
    ]
)

# -- Save --
out_path = Path(__file__).parent / 'TSG_TSE_MDA_Design_Implementation.docx'
doc.save(str(out_path))
print(f'✅ Document saved to: {out_path}')
