"""Generate TSG Executive Overview as DOCX."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Test Suite Generator (TSG)', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Executive Overview')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

doc.add_heading('What It Does', level=1)
doc.add_paragraph(
    'TSG is an AI-powered test suite generation platform that automatically creates '
    'production-ready test cases for TMO/Spectrum Mobile provisioning features. '
    'It replaces weeks of manual test case writing with automated, intelligent generation in minutes.')

doc.add_heading('How It Works', level=1)
doc.add_paragraph(
    'The platform ingests data from three sources automatically, processes it through '
    'an intelligent engine, and outputs production-ready test suites.')

t = doc.add_table(rows=4, cols=2)
t.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate([
    ('Input Source', 'What It Provides'),
    ('Chalk (Design Specs)', 'Feature scenarios, validation rules, transaction flows, CDR derivation rules'),
    ('Jira (Requirements)', 'Acceptance criteria, subtask tables (YL state change matrix), comments, attachments'),
    ('NBOP UI Crawler', 'Real portal menu paths, field names, button labels for UI test steps'),
]):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_paragraph('')
doc.add_paragraph('Output: Production-ready Excel test suites with Summary, Description, Preconditions, Test Steps, and Expected Results.')

doc.add_heading('Key Numbers (PI-52 and PI-53)', level=1)
t2 = doc.add_table(rows=7, cols=2)
t2.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate([
    ('Metric', 'Value'), ('Features covered', '54'), ('Test cases generated', '1,789'),
    ('Test steps generated', '7,233'), ('Quality audit pass rate', '100%'),
    ('Chalk scenario alignment', '97.5%'), ('Avg generation time/feature', '~30 seconds'),
]):
    t2.cell(i, 0).text = k
    t2.cell(i, 1).text = v

doc.add_heading('Intelligence Layers', level=1)

doc.add_heading('Integration Contract', level=2)
doc.add_paragraph(
    '24 registered operations with explicit rules for which external systems '
    '(Syniverse, ITMBO, EMM, APOLLO_NE) each operation calls and which it must NOT call.')

doc.add_heading('Feature Classification', level=2)
doc.add_paragraph(
    'Automatically detects: API, UI, Hybrid, CDR/Mediation, Notification, '
    'Inquiry, Sync Subscriber, and Batch. Each type gets domain-specific step templates.')

doc.add_heading('15+ Domain-Specific Step Templates', level=2)
for b in [
    'Swap MDN: 19-step pipeline with PSIM vs ESIM Syniverse differentiation',
    'Sync Subscriber: Dynamic steps from Jira subtask tables (TMO Status x NSL Status x Syniverse Action)',
    'Order Inquiry: Response payload validation with TMO vs VZW differentiation',
    'CDR/ILD Processing: Mediation pipeline with PRR file verification',
    'Hotline/Remove Hotline: Explicit dual assertion (what happens + what does NOT happen)',
    'Kafka/BI Events: Century Report EVENT_MESSAGES table verification',
    'Activation: Syniverse CreateSubscriber with NBOP Line Summary verification',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Jira Subtask Table Intelligence', level=2)
doc.add_paragraph(
    'When Jira subtasks contain structured tables (e.g., YL state change matrix), '
    'the engine parses them and generates TCs with steps derived directly from the table data.')

doc.add_heading('UI Verification Mirror', level=2)
doc.add_paragraph(
    'For every API operation with an NBOP counterpart, an additional UI verification TC '
    'is auto-generated with real navigation paths from the NBOP crawler.')

doc.add_heading('Quality Assurance', level=1)
doc.add_paragraph('Every generated test suite passes through a multi-layer quality gate:')
for item in [
    'Surface quality: No trailing dots, empty fields, or non-standard categories',
    'Cross-contamination: CDR features never get API steps; inquiry features never get Century Report steps',
    'Chalk alignment: 97.5% of Chalk scenarios have matching TC steps',
    'Semantic alignment: Steps match the feature type',
    'Precondition relevance: Sync scenarios get state-specific preconditions',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Test Suite Executor (TSE)', level=1)
doc.add_paragraph('Generated test suites feed directly into TSE for automated execution:')
for item in [
    'API steps: OAuth + cURL execution with response validation',
    'UI steps: Playwright-driven NBOP portal automation with screenshots',
    'Century Report steps: Automated report download and validation',
    'Evidence capture: Screenshots, API responses, and DOCX evidence documents per TC',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Architecture', level=1)
t3 = doc.add_table(rows=7, cols=2)
t3.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate([
    ('Component', 'Technology'), ('Language', 'Python 3.13'), ('UI', 'Streamlit'),
    ('Browser Automation', 'Playwright (Edge/Chrome)'), ('Database', 'SQLite (local cache)'),
    ('Source Control', 'GitHub'), ('AI/LLM', 'None - all rule-based and deterministic'),
]):
    t3.cell(i, 0).text = k
    t3.cell(i, 1).text = v

out = 'TestSuiteGenerator/outputs/TSG_Executive_Overview.docx'
doc.save(out)
print('Saved: %s' % out)
