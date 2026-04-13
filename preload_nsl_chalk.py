"""
preload_nsl_chalk.py — Scrape NSL Services Chalk page and store in SQLite + DOCX.

Navigates to: https://chalk.charter.com/spaces/MDA/pages/1692266708/NSL+Services
1. Fetches the main right-hand table (Service Id, Service Name, API Type, etc.)
2. Clicks ONLY the child items under "NSL Services" in the left sidebar tree
   (activate-subscriber, whl-activate-subscriber, add-wearable, select-mdn, etc.)
3. For each child item — scrapes the right panel (tables, text, links, sub-items)
4. Stores everything in SQLite table NSL_VZ_Chalk
5. Generates a DOCX document with all fetched data

Usage:  python preload_nsl_chalk.py
        python preload_nsl_chalk.py --force          (re-fetch even if DB has data)
        python preload_nsl_chalk.py --headless false  (visible browser for debugging)
"""
import sys, os, time, argparse, json, re, sqlite3
from pathlib import Path
from datetime import datetime

sys.path.insert(0, os.path.dirname(__file__))

from playwright.sync_api import sync_playwright
from modules.config import get_browser_channel, PAGE_LOAD_TIMEOUT_MS, NETWORK_IDLE_TIMEOUT_MS, ROOT

NSL_URL = 'https://chalk.charter.com/spaces/MDA/pages/1692266708/NSL+Services'
NSL_PAGE_ID = '1692266708'
DB_PATH = ROOT / 'tsg_cache.db'
DOCX_PATH = ROOT / 'outputs' / 'NSL_VZ_Chalk.docx'


def log(msg):
    print(msg, flush=True)


def _wait(page, timeout=NETWORK_IDLE_TIMEOUT_MS):
    try:
        page.wait_for_load_state('networkidle', timeout=timeout)
    except Exception:
        pass


# ================================================================
# DATABASE
# ================================================================

def _conn():
    c = sqlite3.connect(str(DB_PATH), timeout=10)
    c.execute('PRAGMA journal_mode=WAL')
    c.row_factory = sqlite3.Row
    return c


def init_nsl_table():
    """Create the NSL_VZ_Chalk table if it doesn't exist."""
    c = _conn()
    c.executescript('''
        CREATE TABLE IF NOT EXISTS NSL_VZ_Chalk (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            section_name    TEXT NOT NULL,
            section_url     TEXT,
            parent_section  TEXT DEFAULT 'NSL Services',
            content_type    TEXT DEFAULT 'table',
            table_data_json TEXT,
            raw_text        TEXT,
            links_json      TEXT,
            sub_sections_json TEXT,
            last_fetched    TEXT,
            UNIQUE(section_name, parent_section)
        );
    ''')
    c.commit()
    c.close()
    log('  [DB] NSL_VZ_Chalk table ready.')


def save_nsl_section(section_name, section_url, parent_section, content_type,
                     table_data, raw_text, links, sub_sections):
    c = _conn()
    now = datetime.now().isoformat()
    c.execute(
        'INSERT OR REPLACE INTO NSL_VZ_Chalk '
        '(section_name, section_url, parent_section, content_type, '
        'table_data_json, raw_text, links_json, sub_sections_json, last_fetched) '
        'VALUES (?,?,?,?,?,?,?,?,?)',
        (
            section_name, section_url, parent_section, content_type,
            json.dumps(table_data, ensure_ascii=False) if table_data else None,
            raw_text,
            json.dumps(links, ensure_ascii=False) if links else None,
            json.dumps(sub_sections, ensure_ascii=False) if sub_sections else None,
            now
        )
    )
    c.commit()
    c.close()


def get_nsl_row_count():
    c = _conn()
    try:
        r = c.execute('SELECT COUNT(*) FROM NSL_VZ_Chalk').fetchone()[0]
    except Exception:
        r = 0
    c.close()
    return r


def load_all_nsl_data():
    c = _conn()
    rows = c.execute('SELECT * FROM NSL_VZ_Chalk ORDER BY id').fetchall()
    c.close()
    return [dict(r) for r in rows]


# ================================================================
# SCRAPING HELPERS
# ================================================================

def scrape_tables(page):
    """Extract all tables from the main content area."""
    return page.evaluate(r"""() => {
        const tables = [];
        const contentArea = document.querySelector('#main-content') ||
                            document.querySelector('.wiki-content') ||
                            document.querySelector('[data-testid="page-content"]') ||
                            document.body;
        contentArea.querySelectorAll('table').forEach(tbl => {
            const headers = [];
            tbl.querySelectorAll('thead th, tr:first-child th').forEach(th => {
                headers.push(th.innerText.trim());
            });
            if (headers.length === 0) {
                const firstRow = tbl.querySelector('tr');
                if (firstRow) {
                    firstRow.querySelectorAll('td, th').forEach(cell => {
                        headers.push(cell.innerText.trim());
                    });
                }
            }
            const rows = [];
            const allRows = tbl.querySelectorAll('tr');
            const startIdx = headers.length > 0 ? 1 : 0;
            for (let i = startIdx; i < allRows.length; i++) {
                const cells = allRows[i].querySelectorAll('td, th');
                const row = {};
                cells.forEach((cell, ci) => {
                    const key = ci < headers.length ? headers[ci] : ('col_' + ci);
                    row[key] = cell.innerText.trim();
                });
                if (Object.values(row).some(v => v)) rows.push(row);
            }
            if (rows.length > 0) tables.push({headers: headers, rows: rows});
        });
        return tables;
    }""")


def scrape_raw_text(page):
    """Get all visible text from the main content area."""
    return page.evaluate(r"""() => {
        const el = document.querySelector('#main-content') ||
                   document.querySelector('.wiki-content') ||
                   document.querySelector('[data-testid="page-content"]') ||
                   document.body;
        return el ? el.innerText.trim() : '';
    }""")


def scrape_links(page):
    """Extract all links from the main content area."""
    return page.evaluate(r"""() => {
        const links = [];
        const el = document.querySelector('#main-content') ||
                   document.querySelector('.wiki-content') ||
                   document.querySelector('[data-testid="page-content"]') ||
                   document.body;
        el.querySelectorAll('a[href]').forEach(a => {
            const text = a.innerText.trim();
            const href = a.href;
            if (text && href && !href.startsWith('javascript:')) {
                links.push({text: text, href: href});
            }
        });
        return links;
    }""")


def expand_nsl_tree(page):
    """Expand ONLY the NSL Services node and its children in the sidebar tree."""
    for attempt in range(5):
        expanded = page.evaluate(r"""() => {
            let c = 0;
            const NSL_ID = '1692266708';

            // Find the NSL Services anchor by page ID
            let nslAnchor = null;
            document.querySelectorAll('a[href]').forEach(a => {
                if (a.href && a.href.indexOf(NSL_ID) !== -1) nslAnchor = a;
            });
            if (!nslAnchor) return 0;

            // Walk up to the owning <li>
            let nslLi = nslAnchor.closest('li');
            if (!nslLi) return 0;

            // Expand toggle buttons/links ONLY within this <li> subtree
            nslLi.querySelectorAll(
                'button[aria-expanded="false"], ' +
                'a.plugin_pagetree_childtoggle_haschildren'
            ).forEach(btn => {
                try { btn.click(); c++; } catch(e) {}
            });

            return c;
        }""")
        if expanded:
            log('    Expanded %d NSL tree nodes (pass %d)' % (expanded, attempt + 1))
            time.sleep(2)
        else:
            break
    time.sleep(1)


def discover_nsl_children(page):
    """
    Find ONLY the direct child pages under 'NSL Services' in the left sidebar.
    Anchors to the NSL Services <li> node, then collects links from its
    immediate children container — never leaves that subtree.
    """
    items = page.evaluate(r"""() => {
        const items = [];
        const seen = new Set();
        const NSL_ID = '1692266708';

        // ---------- Strategy 1: Anchor to NSL Services <li> in page-tree ----------
        let nslAnchor = null;
        document.querySelectorAll('a[href]').forEach(a => {
            if (a.href && a.href.indexOf(NSL_ID) !== -1) nslAnchor = a;
        });

        if (nslAnchor) {
            let nslLi = nslAnchor.closest('li');
            if (nslLi) {
                // Get the children container (the <ul> or div right under this <li>)
                let childContainer = nslLi.querySelector(
                    ':scope > ul, ' +
                    ':scope > div > ul, ' +
                    ':scope > .plugin_pagetree_children_container, ' +
                    ':scope > div.plugin_pagetree_children_container'
                );
                let root = childContainer || nslLi;

                // Collect ONLY direct child <li> links (one level)
                let childLis = root.querySelectorAll(':scope > li');
                if (childLis.length > 0) {
                    childLis.forEach(li => {
                        let a = li.querySelector('a[href*="/spaces/MDA/pages/"]');
                        if (!a) a = li.querySelector('a[href]');
                        if (a) {
                            const text = a.innerText.trim();
                            const href = a.href;
                            if (text && href && !seen.has(href) &&
                                href.indexOf(NSL_ID + '/NSL') === -1 &&
                                text !== 'NSL Services') {
                                seen.add(href);
                                items.push({text: text, href: href});
                            }
                        }
                    });
                } else {
                    // Fallback: get all links inside the children container
                    root.querySelectorAll('a[href*="/spaces/MDA/pages/"]').forEach(a => {
                        const text = a.innerText.trim();
                        const href = a.href;
                        if (text && href && !seen.has(href) &&
                            href.indexOf(NSL_ID) === -1 &&
                            text !== 'NSL Services') {
                            seen.add(href);
                            items.push({text: text, href: href});
                        }
                    });
                }
            }
        }

        // ---------- Strategy 2: Confluence Server plugin_pagetree ----------
        if (items.length === 0) {
            let currentNode = document.querySelector('.plugin_pagetree_current') ||
                              document.querySelector('[data-pageid="' + NSL_ID + '"]');
            if (currentNode) {
                let parentLi = currentNode.closest('li') || currentNode.parentElement;
                let childContainer = parentLi ?
                    parentLi.querySelector('.plugin_pagetree_children_container') : null;
                if (childContainer) {
                    childContainer.querySelectorAll(':scope > div > a, :scope a[href*="/spaces/MDA/pages/"]').forEach(a => {
                        const text = a.innerText.trim();
                        const href = a.href;
                        if (text && href && !seen.has(href) && href.indexOf(NSL_ID) === -1) {
                            seen.add(href);
                            items.push({text: text, href: href});
                        }
                    });
                }
            }
        }

        // ---------- Strategy 3: Depth-based walk ----------
        // Walk the sidebar links in DOM order. Once we hit the NSL Services link,
        // collect subsequent links that are deeper (children). Stop when we hit
        // a link at the same or shallower depth (sibling/parent).
        if (items.length === 0) {
            const allLinks = document.querySelectorAll(
                '.ia-splitter-left a[href*="/spaces/MDA/pages/"], ' +
                'nav a[href*="/spaces/MDA/pages/"]'
            );
            let foundNsl = false;
            let nslDepth = -1;

            function liDepth(el) {
                let d = 0;
                while (el) { if (el.tagName === 'LI') d++; el = el.parentElement; }
                return d;
            }

            allLinks.forEach(a => {
                const text = a.innerText.trim();
                const href = a.href;
                if (!foundNsl) {
                    if (href.indexOf(NSL_ID) !== -1) {
                        foundNsl = true;
                        nslDepth = liDepth(a);
                    }
                    return;
                }
                // We are past NSL Services — check depth
                const d = liDepth(a);
                if (d > nslDepth && text && href && !seen.has(href)) {
                    seen.add(href);
                    items.push({text: text, href: href});
                } else if (d <= nslDepth) {
                    // Exited the NSL subtree — stop collecting
                    foundNsl = false;
                }
            });
        }

        return items;
    }""")

    return items


# ================================================================
# DOCX GENERATION
# ================================================================

def generate_docx(all_data):
    """Generate a DOCX document with all scraped NSL data."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        log('[DOCX] python-docx not installed. Run: pip install python-docx')
        log('[DOCX] Skipping DOCX generation.')
        return None

    doc = Document()
    doc.add_heading('NSL VZ Chalk \u2014 Service Data Export', level=0)
    doc.add_paragraph('Generated: %s' % datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    doc.add_paragraph('Source: %s' % NSL_URL)
    doc.add_paragraph('Total sections scraped: %d' % len(all_data))
    doc.add_paragraph('')

    for section in all_data:
        name = section.get('section_name', 'Unknown')
        parent = section.get('parent_section', '')
        url = section.get('section_url', '')
        raw = section.get('raw_text', '')
        table_json = section.get('table_data_json')
        links_json = section.get('links_json')
        subs_json = section.get('sub_sections_json')

        heading_text = '%s > %s' % (parent, name) if parent and parent != name else name
        doc.add_heading(heading_text, level=1)

        if url:
            doc.add_paragraph('URL: %s' % url)

        # Tables
        tables = []
        if table_json:
            try:
                tables = json.loads(table_json) if isinstance(table_json, str) else table_json
            except Exception:
                pass

        if tables:
            for ti, tbl in enumerate(tables):
                if ti > 0:
                    doc.add_paragraph('')
                headers = tbl.get('headers', [])
                rows = tbl.get('rows', [])
                if not rows:
                    continue
                cols = headers if headers else list(rows[0].keys()) if rows else []
                if not cols:
                    continue

                doc.add_paragraph('Table %d (%d rows):' % (ti + 1, len(rows)))
                tbl_obj = doc.add_table(rows=1 + len(rows), cols=len(cols))
                tbl_obj.style = 'Table Grid'

                for ci, col in enumerate(cols):
                    cell = tbl_obj.rows[0].cells[ci]
                    cell.text = str(col)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

                for ri, row in enumerate(rows):
                    for ci, col in enumerate(cols):
                        tbl_obj.rows[ri + 1].cells[ci].text = str(row.get(col, ''))

        # Links
        links = []
        if links_json:
            try:
                links = json.loads(links_json) if isinstance(links_json, str) else links_json
            except Exception:
                pass

        if links:
            doc.add_heading('Links', level=2)
            for lnk in links:
                doc.add_paragraph('\u2022 %s \u2014 %s' % (lnk.get('text', ''), lnk.get('href', '')))

        # Sub-sections
        subs = []
        if subs_json:
            try:
                subs = json.loads(subs_json) if isinstance(subs_json, str) else subs_json
            except Exception:
                pass

        if subs:
            doc.add_heading('Sub-sections', level=2)
            for sub in subs:
                doc.add_paragraph('\u2022 %s \u2014 %s' % (sub.get('text', ''), sub.get('href', '')))

        # Raw text fallback (when no tables)
        if raw and not tables:
            doc.add_heading('Content', level=2)
            text_to_add = raw[:5000]
            if len(raw) > 5000:
                text_to_add += '\n... [truncated \u2014 see DB for full text]'
            doc.add_paragraph(text_to_add)

        doc.add_page_break()

    DOCX_PATH.parent.mkdir(exist_ok=True)
    doc.save(str(DOCX_PATH))
    log('  [DOCX] Saved: %s' % DOCX_PATH)
    return DOCX_PATH


# ================================================================
# MAIN SCRAPER
# ================================================================

def scrape_nsl(force=False, headless=True):
    init_nsl_table()

    if not force and get_nsl_row_count() > 0:
        count = get_nsl_row_count()
        log('DB already has %d NSL sections. Use --force to re-fetch.' % count)
        log('Generating DOCX from existing data...')
        all_data = load_all_nsl_data()
        generate_docx(all_data)
        return

    t0 = time.time()
    log('=' * 60)
    log('  NSL VZ Chalk Scraper')
    log('  Target: %s' % NSL_URL)
    log('  DB:     %s' % DB_PATH)
    log('=' * 60)

    log('\nLaunching browser...')
    pw = sync_playwright().start()
    br = pw.chromium.launch(headless=headless, channel=get_browser_channel())
    ctx = br.new_context(viewport={'width': 1920, 'height': 1080})
    page = ctx.new_page()

    # ---- Step 1: Load the NSL Services page ----
    log('\n[1/3] Loading NSL Services page...')
    page.goto(NSL_URL, wait_until='domcontentloaded', timeout=PAGE_LOAD_TIMEOUT_MS)
    _wait(page)
    time.sleep(5)
    log('  Page title: %s' % page.title()[:80])

    # Verify we are on the right page
    current_url = page.url
    if NSL_PAGE_ID not in current_url:
        log('  WARNING: Current URL does not contain NSL page ID!')
        log('  Current URL: %s' % current_url)
        log('  Expected page ID: %s' % NSL_PAGE_ID)

    # ---- Step 2: Scrape the main right-hand table ----
    log('\n[2/3] Scraping main page table (right panel)...')
    main_tables = scrape_tables(page)
    main_text = scrape_raw_text(page)
    main_links = scrape_links(page)

    log('  Found %d table(s) on main page' % len(main_tables))
    for i, t in enumerate(main_tables):
        log('    Table %d: %d rows, headers=%s' % (i + 1, len(t.get('rows', [])), t.get('headers', [])[:5]))

    save_nsl_section(
        section_name='NSL Services (Main)',
        section_url=NSL_URL,
        parent_section='NSL Services',
        content_type='table',
        table_data=main_tables,
        raw_text=main_text[:50000],
        links=main_links,
        sub_sections=None
    )
    log('  Saved main page data.')

    # ---- Step 3: Expand NSL tree and discover ONLY its children ----
    log('\n[3/3] Discovering child items under NSL Services...')
    expand_nsl_tree(page)
    sidebar_items = discover_nsl_children(page)
    log('  Found %d child items under NSL Services' % len(sidebar_items))

    for item in sidebar_items:
        log('    - %s  (%s)' % (item.get('text', ''), item.get('href', '')[:80]))

    if not sidebar_items:
        log('  WARNING: No child items found! The page tree may not have loaded.')
        log('  Try running with --headless false to debug.')

    total_scraped = 1  # main page already done
    failed = 0

    for idx, item in enumerate(sidebar_items, 1):
        item_name = item.get('text', '').strip()
        item_url = item.get('href', '')

        if not item_name or not item_url:
            continue

        # Safety check: only follow links within /spaces/MDA/pages/
        if '/spaces/MDA/pages/' not in item_url:
            log('  [%d/%d] SKIP (not an MDA page): %s' % (idx, len(sidebar_items), item_name))
            continue

        log('\n  [%d/%d] %s' % (idx, len(sidebar_items), item_name))
        log('    URL: %s' % item_url[:120])

        try:
            page.goto(item_url, wait_until='domcontentloaded', timeout=PAGE_LOAD_TIMEOUT_MS)
            _wait(page)
            time.sleep(3)

            # Verify we navigated to the right page (not redirected elsewhere)
            actual_url = page.url
            log('    Actual URL: %s' % actual_url[:120])

            child_tables = scrape_tables(page)
            child_text = scrape_raw_text(page)
            child_links = scrape_links(page)

            # Check for sub-pages within this child page
            child_sub_items = page.evaluate(r"""() => {
                const items = [];
                const seen = new Set();
                document.querySelectorAll(
                    '.childpages-macro a, .children-show-if a, ' +
                    '.plugin_pagetree_children_container a'
                ).forEach(a => {
                    const text = a.innerText.trim();
                    const href = a.href;
                    if (text && href && !seen.has(href)) {
                        seen.add(href);
                        items.push({text: text, href: href});
                    }
                });
                return items;
            }""")

            content_type = 'table' if child_tables else 'text'
            table_count = len(child_tables)
            row_count = sum(len(t.get('rows', [])) for t in child_tables)

            save_nsl_section(
                section_name=item_name,
                section_url=item_url,
                parent_section='NSL Services',
                content_type=content_type,
                table_data=child_tables,
                raw_text=child_text[:50000],
                links=child_links,
                sub_sections=child_sub_items if child_sub_items else None
            )

            total_scraped += 1
            log('    OK \u2014 %d table(s), %d rows, %d links, %d sub-items' % (
                table_count, row_count, len(child_links), len(child_sub_items)))

        except Exception as e:
            failed += 1
            log('    FAILED: %s' % str(e)[:100])

    # Cleanup browser
    ctx.close()
    br.close()
    pw.stop()

    # ---- Generate DOCX ----
    log('\nGenerating DOCX report...')
    all_data = load_all_nsl_data()
    generate_docx(all_data)

    # Summary
    elapsed = time.time() - t0
    m, s = divmod(int(elapsed), 60)
    log('\n' + '=' * 60)
    log('  NSL SCRAPE COMPLETE')
    log('  Sections scraped: %d' % total_scraped)
    log('  Failed:           %d' % failed)
    log('  DB rows:          %d' % get_nsl_row_count())
    log('  DB path:          %s' % DB_PATH)
    log('  DOCX path:        %s' % DOCX_PATH)
    log('  Time:             %dm %ds' % (m, s))
    log('=' * 60)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Scrape NSL Services from Chalk into DB + DOCX')
    parser.add_argument('--force', action='store_true', help='Re-fetch even if DB has data')
    parser.add_argument('--headless', type=str, default='true',
                        help='Run browser headless (true/false). Default: true')
    args = parser.parse_args()
    headless = args.headless.lower() != 'false'
    scrape_nsl(force=args.force, headless=headless)
