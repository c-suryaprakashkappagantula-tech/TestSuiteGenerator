"""
doc_generator.py — Generate Feature Objective & Summary document (.docx)
Auto-generated alongside the test suite Excel, pulling from Jira + Chalk + Suite data.
"""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .config import OUTPUTS


def generate_feature_doc(suite, jira, chalk, log=print) -> Path:
    """Generate a Feature Objective & Test Summary document.
    Returns the output path."""
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Title ──
    title = doc.add_heading('Feature Objective & Test Summary', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # ── Feature Overview Table ──
    doc.add_heading('1. Feature Overview', level=1)
    _add_info_table(doc, [
        ('Feature ID', suite.feature_id),
        ('Title', suite.feature_title),
        ('PI', suite.pi or 'N/A'),
        ('Status', jira.status if jira else 'N/A'),
        ('Priority', jira.priority if jira else 'N/A'),
        ('Assignee', jira.assignee if jira else 'N/A'),
        ('Channel', ', '.join(suite.channel) if isinstance(suite.channel, list) else suite.channel),
        ('Labels', ', '.join(jira.labels) if jira and jira.labels else 'None'),
        ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M')),
    ])

    # ── Objective ──
    doc.add_heading('2. Objective', level=1)
    _obj_text = 'Validate the %s feature for TMO MVNO provisioning. ' % suite.feature_title
    if chalk and chalk.scope:
        _obj_text += chalk.scope[:500]
    elif jira and jira.description:
        _obj_text += jira.description[:500]
    doc.add_paragraph(_obj_text)

    # ── Scope & Rules ──
    if chalk and (chalk.scope or chalk.rules):
        doc.add_heading('3. Scope & Business Rules', level=1)
        if chalk.scope:
            doc.add_paragraph(chalk.scope[:1000])
        if chalk.rules:
            doc.add_paragraph(chalk.rules[:1000])

    # ── Acceptance Criteria ──
    if suite.acceptance_criteria:
        doc.add_heading('4. Acceptance Criteria', level=1)
        for i, ac in enumerate(suite.acceptance_criteria, 1):
            doc.add_paragraph('%d. %s' % (i, ac), style='List Number')

    # ── Test Coverage Summary ──
    doc.add_heading('5. Test Coverage Summary', level=1)

    # Category breakdown
    cats = {}
    for tc in suite.test_cases:
        cats.setdefault(tc.category, []).append(tc)

    _add_summary_table(doc, [
        ('Total Test Cases', str(len(suite.test_cases))),
        ('Total Test Steps', str(sum(len(tc.steps) for tc in suite.test_cases))),
    ] + [('%s Scenarios' % cat, str(len(tcs))) for cat, tcs in cats.items()])

    # Priority breakdown
    pris = {}
    for tc in suite.test_cases:
        pri = getattr(tc, '_priority', 'P3')
        pris.setdefault(pri, []).append(tc)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Priority Distribution: ').bold = True
    for pri in ['P1', 'P2', 'P3']:
        if pri in pris:
            p.add_run('%s=%d  ' % (pri, len(pris[pri])))

    # ── Scenario List ──
    doc.add_heading('6. Test Scenarios', level=1)

    for cat, tcs in cats.items():
        doc.add_heading(cat, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        hdr[0].text = 'TC#'
        hdr[1].text = 'Scenario'
        hdr[2].text = 'Priority'
        for cell in hdr:
            for p in cell.paragraphs:
                p.runs[0].bold = True if p.runs else False

        for tc in tcs:
            row = table.add_row().cells
            row[0].text = 'TC%s' % tc.sno.zfill(2)
            # Clean summary for doc
            _name = tc.summary
            _name = _name.split('_', 2)[-1] if '_' in _name else _name  # strip TC##_FEAT_
            if _name.startswith(suite.feature_id):
                _name = _name[len(suite.feature_id):].lstrip('_ ')
            row[1].text = _name[:120]
            row[2].text = getattr(tc, '_priority', 'P3')

    # ── Data Sources ──
    doc.add_heading('7. Data Sources', level=1)
    for src in suite.data_sources:
        doc.add_paragraph(src, style='List Bullet')

    # ── Linked Issues ──
    if jira and jira.linked_issues:
        doc.add_heading('8. Linked Issues', level=1)
        for link in jira.linked_issues:
            doc.add_paragraph('%s — %s' % (link.get('key', ''), link.get('summary', '')[:80]), style='List Bullet')

    # ── Warnings / Open Items ──
    if suite.warnings or suite.open_items:
        doc.add_heading('9. Warnings & Open Items', level=1)
        for w in suite.warnings:
            doc.add_paragraph(w, style='List Bullet')
        for item in suite.open_items[:10]:
            doc.add_paragraph(item, style='List Bullet')

    # ── Save ──
    _ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    out_name = '%s_Feature_Summary_%s.docx' % (suite.feature_id, _ts)
    out_path = OUTPUTS / out_name
    doc.save(str(out_path))
    log('[DOC-GEN] Feature summary saved: %s' % out_name)
    return out_path


def _add_info_table(doc, rows):
    """Add a 2-column info table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value or '')
        for p in table.rows[i].cells[0].paragraphs:
            if p.runs:
                p.runs[0].bold = True


def _add_summary_table(doc, rows):
    """Add a 2-column summary stats table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
