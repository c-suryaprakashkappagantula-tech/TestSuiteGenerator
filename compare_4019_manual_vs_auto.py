"""
Compare Manual Test Suite vs TSG Dashboard Generated Test Suite for MWTGPROV-4019 (Sync Key Info).
Manual:  .xlsx from QMetry/manual export
Auto:    .docx Feature Summary from TSG Dashboard
"""
import openpyxl
import re
from docx import Document
import sys

MANUAL = r'C:\Users\P3314665\Downloads\testcases_4019_1777349025192_oNeUiEj5Kk.xlsx'
AUTO   = r'C:\Users\P3314665\Documents\mcp-jira-server\TestSuiteGenerator\outputs\MWTGPROV-4019_Feature_Summary_20260428_093125.docx'

# Also try to find the latest .xlsx auto-generated file for 4019
import glob, os
AUTO_XLSX_PATTERN = r'C:\Users\P3314665\Documents\mcp-jira-server\TestSuiteGenerator\outputs\MWTGPROV-4019_*_20260428_093124.xlsx'
auto_xlsx_files = glob.glob(AUTO_XLSX_PATTERN)
if not auto_xlsx_files:
    AUTO_XLSX_PATTERN = r'C:\Users\P3314665\Documents\mcp-jira-server\TestSuiteGenerator\outputs\MWTGPROV-4019_*_20260428*.xlsx'
    auto_xlsx_files = glob.glob(AUTO_XLSX_PATTERN)
if not auto_xlsx_files:
    # Get the latest 4019 xlsx
    AUTO_XLSX_PATTERN = r'C:\Users\P3314665\Documents\mcp-jira-server\TestSuiteGenerator\outputs\MWTGPROV-4019_*.xlsx'
    auto_xlsx_files = sorted(glob.glob(AUTO_XLSX_PATTERN))

AUTO_XLSX = auto_xlsx_files[-1] if auto_xlsx_files else None

# ─── EXTRACT MANUAL TCs FROM XLSX ───────────────────────────────────────────
def extract_tcs_from_xlsx(path, label):
    wb = openpyxl.load_workbook(path, data_only=True)
    tcs = []
    for sname in wb.sheetnames:
        ws = wb[sname]
        rows = list(ws.iter_rows(min_row=1, values_only=True))
        if not rows:
            continue
        # Find header row
        header_idx = -1
        for i, row in enumerate(rows):
            row_str = ' '.join(str(c or '').lower() for c in row)
            if any(kw in row_str for kw in ['s.no', 'summary', 'test case', 'sno', 'tc id', 'tc name', 'test name', 'test_case_name']):
                header_idx = i
                break
        if header_idx < 0:
            header_idx = 0 if len(rows) > 2 else -1
        if header_idx < 0:
            continue

        headers = [str(c or '').strip().lower() for c in rows[header_idx]]
        col_map = {}
        for ci, h in enumerate(headers):
            if any(k in h for k in ['s.no', 'sno', 'tc id', 'tc_id', 'issue key']):
                col_map.setdefault('sno', ci)
            elif h == 'id':
                col_map.setdefault('sno', ci)
            if any(k in h for k in ['summary', 'test case name', 'test_case_name', 'tc name', 'test name', 'scenario', 'test scenario']):
                col_map.setdefault('summary', ci)
            if any(k in h for k in ['description', 'desc', 'test_case_description']):
                col_map.setdefault('description', ci)
            if any(k in h for k in ['precondition', 'pre-condition', 'pre condition']):
                col_map.setdefault('preconditions', ci)
            if ('step' in h and any(k in h for k in ['summary', 'action', 'detail', 'name'])) or h == 'steps' or h == 'test steps' or h == 'test_steps':
                col_map.setdefault('step_summary', ci)
            if 'expected' in h:
                col_map.setdefault('expected', ci)
            if any(k in h for k in ['category', 'type', 'test type']):
                col_map.setdefault('category', ci)
            if any(k in h for k in ['label', 'tag']):
                col_map.setdefault('labels', ci)
            if any(k in h for k in ['priority']):
                col_map.setdefault('priority', ci)

        current_tc = None
        for row in rows[header_idx+1:]:
            row_vals = list(row)
            if len(row_vals) == 0:
                continue
            
            def safe_get(idx, default=''):
                if idx is not None and idx < len(row_vals):
                    return str(row_vals[idx] or '').strip()
                return default

            sno_val = safe_get(col_map.get('sno'))
            summary_val = safe_get(col_map.get('summary'))

            if sno_val and sno_val not in ('None', ''):
                if current_tc:
                    tcs.append(current_tc)
                current_tc = {
                    'sheet': sname,
                    'sno': sno_val,
                    'summary': summary_val,
                    'description': safe_get(col_map.get('description')),
                    'preconditions': safe_get(col_map.get('preconditions')),
                    'steps': [],
                    'category': safe_get(col_map.get('category')),
                    'priority': safe_get(col_map.get('priority')),
                    'labels': safe_get(col_map.get('labels')),
                }
            elif summary_val and not current_tc:
                current_tc = {
                    'sheet': sname,
                    'sno': sno_val or '?',
                    'summary': summary_val,
                    'description': safe_get(col_map.get('description')),
                    'preconditions': safe_get(col_map.get('preconditions')),
                    'steps': [],
                    'category': '',
                    'priority': '',
                    'labels': '',
                }

            if current_tc:
                step_sum = safe_get(col_map.get('step_summary'))
                step_exp = safe_get(col_map.get('expected'))
                if step_sum or step_exp:
                    current_tc['steps'].append({'summary': step_sum, 'expected': step_exp})

        if current_tc:
            tcs.append(current_tc)
    wb.close()
    return tcs


# ─── EXTRACT AUTO TCs FROM DOCX ─────────────────────────────────────────────
def extract_tcs_from_docx(path):
    doc = Document(path)
    tcs = []
    current_tc = None
    in_steps = False
    feature_info = {}
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.strip())
    
    # Also extract from tables
    table_tcs = []
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue
        headers = [cell.text.strip().lower() for cell in rows[0].cells]
        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            tc = {}
            for i, h in enumerate(headers):
                if i < len(cells):
                    tc[h] = cells[i]
            if tc:
                table_tcs.append(tc)
    
    # Parse paragraphs for TC structure
    i = 0
    while i < len(full_text):
        line = full_text[i]
        
        # Detect TC headers like "TC01:", "Test Case 1:", "TC-01", etc.
        tc_match = re.match(r'^(?:TC[-_]?\s*(\d+)|Test\s*Case\s*[-#]?\s*(\d+))\s*[:\-–]\s*(.*)', line, re.IGNORECASE)
        if tc_match:
            if current_tc:
                tcs.append(current_tc)
            tc_num = tc_match.group(1) or tc_match.group(2)
            tc_name = tc_match.group(3).strip()
            current_tc = {
                'sno': tc_num,
                'summary': tc_name,
                'description': '',
                'preconditions': '',
                'steps': [],
                'expected_results': [],
                'category': '',
                'priority': '',
            }
            in_steps = False
            i += 1
            continue
        
        # Detect section headers within a TC
        if current_tc:
            lower = line.lower()
            if lower.startswith('description:') or lower.startswith('objective:'):
                current_tc['description'] = line.split(':', 1)[1].strip()
            elif lower.startswith('precondition') or lower.startswith('pre-condition'):
                current_tc['preconditions'] = line.split(':', 1)[1].strip() if ':' in line else ''
            elif lower.startswith('step') or lower.startswith('test step'):
                in_steps = True
            elif lower.startswith('expected result') or lower.startswith('expected:'):
                val = line.split(':', 1)[1].strip() if ':' in line else ''
                if val:
                    current_tc['expected_results'].append(val)
                in_steps = False
            elif lower.startswith('category:') or lower.startswith('type:'):
                current_tc['category'] = line.split(':', 1)[1].strip()
            elif lower.startswith('priority:'):
                current_tc['priority'] = line.split(':', 1)[1].strip()
            elif in_steps and line and not line.startswith('─'):
                # Parse step lines like "1. Do something" or "- Do something"
                step_match = re.match(r'^[\d]+[.)]\s*(.*)', line)
                if step_match:
                    current_tc['steps'].append({'summary': step_match.group(1), 'expected': ''})
                elif line.startswith('-') or line.startswith('•'):
                    current_tc['steps'].append({'summary': line.lstrip('-•').strip(), 'expected': ''})
                elif current_tc['steps']:
                    # Continuation of previous step
                    current_tc['steps'][-1]['summary'] += ' ' + line
        
        # Detect feature-level info
        if line.lower().startswith('feature:') or line.lower().startswith('feature id:'):
            feature_info['feature'] = line.split(':', 1)[1].strip()
        elif line.lower().startswith('api:') or line.lower().startswith('endpoint:'):
            feature_info['api'] = line.split(':', 1)[1].strip()
        
        i += 1
    
    if current_tc:
        tcs.append(current_tc)
    
    return tcs, table_tcs, feature_info, full_text


# ─── EXTRACT AUTO TCs FROM XLSX ─────────────────────────────────────────────
def extract_auto_xlsx(path):
    if not path:
        return []
    return extract_tcs_from_xlsx(path, 'Auto')


# ─── MAIN COMPARISON ────────────────────────────────────────────────────────
print('='*100)
print('DETAILED COMPARISON: Manual vs TSG Dashboard Generated')
print('Feature: MWTGPROV-4019 — Sync Key Info')
print('='*100)
print(f'\n  Manual file:  {MANUAL}')
print(f'  Auto DOCX:    {AUTO}')
print(f'  Auto XLSX:    {AUTO_XLSX or "Not found"}')

# Extract manual TCs
try:
    manual_tcs = extract_tcs_from_xlsx(MANUAL, 'Manual')
    print(f'\n  ✓ Manual TCs extracted: {len(manual_tcs)}')
except Exception as e:
    print(f'\n  ✗ Error reading manual file: {e}')
    manual_tcs = []

# Extract auto TCs from DOCX
try:
    auto_tcs_docx, auto_table_tcs, feature_info, docx_text = extract_tcs_from_docx(AUTO)
    print(f'  ✓ Auto DOCX TCs (paragraph): {len(auto_tcs_docx)}')
    print(f'  ✓ Auto DOCX table rows: {len(auto_table_tcs)}')
except Exception as e:
    print(f'  ✗ Error reading auto DOCX: {e}')
    auto_tcs_docx, auto_table_tcs, feature_info, docx_text = [], [], {}, []

# Extract auto TCs from XLSX
try:
    auto_tcs_xlsx = extract_auto_xlsx(AUTO_XLSX)
    print(f'  ✓ Auto XLSX TCs extracted: {len(auto_tcs_xlsx)}')
except Exception as e:
    print(f'  ✗ Error reading auto XLSX: {e}')
    auto_tcs_xlsx = []

# Use the best auto source
auto_tcs = auto_tcs_xlsx if auto_tcs_xlsx else auto_tcs_docx
auto_source = 'XLSX' if auto_tcs_xlsx else 'DOCX'

print(f'\n  Using auto source: {auto_source} ({len(auto_tcs)} TCs)')

# ─── SECTION 1: OVERVIEW ────────────────────────────────────────────────────
print('\n' + '─'*100)
print('1. OVERVIEW')
print('─'*100)
print(f'  Manual:    {len(manual_tcs)} TCs')
print(f'  Auto:      {len(auto_tcs)} TCs')
if feature_info:
    for k, v in feature_info.items():
        print(f'  {k}: {v}')

# Show sheets
if manual_tcs:
    manual_sheets = set(tc['sheet'] for tc in manual_tcs)
    print(f'  Manual sheets: {sorted(manual_sheets)}')
if auto_tcs_xlsx:
    auto_sheets = set(tc['sheet'] for tc in auto_tcs_xlsx)
    print(f'  Auto sheets:   {sorted(auto_sheets)}')

# Show column headers found
if manual_tcs:
    wb = openpyxl.load_workbook(MANUAL, data_only=True)
    for sname in wb.sheetnames:
        ws = wb[sname]
        rows = list(ws.iter_rows(min_row=1, max_row=3, values_only=True))
        print(f'\n  Manual sheet "{sname}" headers:')
        for ri, row in enumerate(rows):
            print(f'    Row {ri}: {[str(c or "")[:30] for c in row]}')
    wb.close()

# ─── SECTION 2: TC NAME COMPARISON ──────────────────────────────────────────
print('\n' + '─'*100)
print('2. TC NAME COMPARISON (side by side)')
print('─'*100)
print(f'\n  {"#":>3} {"MANUAL TC NAMES":<65} | {"#":>3} {"AUTO TC NAMES"}')
print(f'  {"─"*3} {"─"*65} | {"─"*3} {"─"*65}')
max_rows = max(len(manual_tcs), len(auto_tcs))
for i in range(max_rows):
    m_name = manual_tcs[i]['summary'][:63] if i < len(manual_tcs) else ''
    a_name = auto_tcs[i]['summary'][:63] if i < len(auto_tcs) else ''
    m_sno = manual_tcs[i]['sno'][:3] if i < len(manual_tcs) else ''
    a_sno = auto_tcs[i]['sno'][:3] if i < len(auto_tcs) else ''
    print(f'  {m_sno:>3} {m_name:<65} | {a_sno:>3} {a_name}')

# ─── SECTION 3: SCENARIO COVERAGE GAP ANALYSIS ──────────────────────────────
print('\n' + '─'*100)
print('3. SCENARIO COVERAGE GAP ANALYSIS')
print('─'*100)

themes = [
    'sync key', 'key info', 'happy path', 'positive', 'negative', 'error',
    'invalid', 'timeout', 'retry', 'rollback', 'concurrent', 'boundary',
    'e2e', 'end-to-end', 'api', 'rest', 'post', 'get', 'put',
    'authentication', 'auth', 'token', 'header',
    'request', 'response', 'payload', 'json', 'xml',
    'missing', 'empty', 'null', 'blank',
    'duplicate', 'idempoten',
    'downstream', 'upstream', 'integration',
    'notification', 'callback', 'webhook',
    'database', 'db', 'persist',
    'performance', 'load', 'stress',
    'security', 'injection', 'xss',
    'field validation', 'format', 'length',
    'status', 'state', 'transition',
    'subscriber', 'account', 'line', 'mdn',
    'sim', 'esim', 'psim', 'iccid', 'imsi', 'imei',
    'nbop', 'nslnm', 'intg', 'itmbo',
    'tmo', 'mvno',
    'century', 'genesis',
    'kafka', 'event',
]

manual_themes = set()
for tc in manual_tcs:
    t = (tc['summary'] + ' ' + tc['description'] + ' ' + tc.get('preconditions', '')).lower()
    for theme in themes:
        if theme in t:
            manual_themes.add(theme)

auto_themes = set()
for tc in auto_tcs:
    t = (tc['summary'] + ' ' + tc.get('description', '') + ' ' + tc.get('preconditions', '')).lower()
    for theme in themes:
        if theme in t:
            auto_themes.add(theme)

print(f'  Manual themes ({len(manual_themes)}): {sorted(manual_themes)}')
print(f'  Auto themes ({len(auto_themes)}):   {sorted(auto_themes)}')
print(f'\n  ⚠ In Manual ONLY (gaps in auto):  {sorted(manual_themes - auto_themes)}')
print(f'  ✓ In Auto ONLY (extra coverage):  {sorted(auto_themes - manual_themes)}')
print(f'  ✓ Common themes:                  {sorted(manual_themes & auto_themes)}')

# ─── SECTION 4: CATEGORY BREAKDOWN ──────────────────────────────────────────
print('\n' + '─'*100)
print('4. CATEGORY / TEST TYPE BREAKDOWN')
print('─'*100)

def categorize_tc(tc):
    t = (tc['summary'] + ' ' + tc.get('description', '') + ' ' + tc.get('category', '')).lower()
    if any(kw in t for kw in ['negative', 'invalid', 'fail', 'reject', 'error', 'missing', 'empty', 'null', 'unauthorized', 'forbidden', 'bad request']):
        return 'Negative'
    elif any(kw in t for kw in ['e2e', 'end-to-end', 'end to end', 'integration']):
        return 'E2E/Integration'
    elif any(kw in t for kw in ['edge', 'boundary', 'concurrent', 'idempoten', 'duplicate', 'timeout', 'retry']):
        return 'Edge Case'
    elif any(kw in t for kw in ['security', 'auth', 'token', 'injection']):
        return 'Security'
    elif any(kw in t for kw in ['performance', 'load', 'stress']):
        return 'Performance'
    else:
        return 'Happy Path'

m_cats = {}
for tc in manual_tcs:
    cat = categorize_tc(tc)
    m_cats[cat] = m_cats.get(cat, 0) + 1

a_cats = {}
for tc in auto_tcs:
    cat = categorize_tc(tc)
    a_cats[cat] = a_cats.get(cat, 0) + 1

all_cats = sorted(set(list(m_cats.keys()) + list(a_cats.keys())))
print(f'    {"Category":<25} {"Manual":>10} {"Auto":>10} {"Delta":>10}')
print(f'    {"─"*25} {"─"*10} {"─"*10} {"─"*10}')
for cat in all_cats:
    m = m_cats.get(cat, 0)
    a = a_cats.get(cat, 0)
    delta = a - m
    sign = '+' if delta > 0 else ''
    print(f'    {cat:<25} {m:>10} {a:>10} {sign}{delta:>9}')
print(f'    {"─"*25} {"─"*10} {"─"*10} {"─"*10}')
print(f'    {"TOTAL":<25} {len(manual_tcs):>10} {len(auto_tcs):>10} {len(auto_tcs)-len(manual_tcs):>+10}')

# ─── SECTION 5: STEP QUALITY ────────────────────────────────────────────────
print('\n' + '─'*100)
print('5. STEP QUALITY COMPARISON')
print('─'*100)
m_total_steps = sum(len(tc['steps']) for tc in manual_tcs)
a_total_steps = sum(len(tc['steps']) for tc in auto_tcs)
m_avg_steps = m_total_steps / max(len(manual_tcs), 1)
a_avg_steps = a_total_steps / max(len(auto_tcs), 1)

m_empty_steps = sum(1 for tc in manual_tcs if len(tc['steps']) == 0)
a_empty_steps = sum(1 for tc in auto_tcs if len(tc['steps']) == 0)

print(f'    {"Metric":<35} {"Manual":>10} {"Auto":>10}')
print(f'    {"─"*35} {"─"*10} {"─"*10}')
print(f'    {"Total TCs":<35} {len(manual_tcs):>10} {len(auto_tcs):>10}')
print(f'    {"Total steps":<35} {m_total_steps:>10} {a_total_steps:>10}')
print(f'    {"Avg steps per TC":<35} {m_avg_steps:>10.1f} {a_avg_steps:>10.1f}')
print(f'    {"TCs with 0 steps":<35} {m_empty_steps:>10} {a_empty_steps:>10}')

# Step detail for first few TCs
for label, tcs_list in [('MANUAL', manual_tcs), ('AUTO', auto_tcs)]:
    print(f'\n  {label} — First 3 TCs detail:')
    for tc in tcs_list[:3]:
        print(f'    [{tc["sno"]}] {tc["summary"][:80]}')
        if tc.get('description'):
            print(f'        Desc: {tc["description"][:100]}')
        if tc.get('preconditions'):
            print(f'        Pre:  {tc["preconditions"][:100]}')
        for si, s in enumerate(tc['steps'][:5]):
            print(f'        Step {si+1}: {s["summary"][:90]}')
            if s.get('expected'):
                print(f'          Exp: {s["expected"][:90]}')
        if len(tc['steps']) > 5:
            print(f'        ... +{len(tc["steps"])-5} more steps')

# ─── SECTION 6: DESCRIPTION QUALITY ─────────────────────────────────────────
print('\n' + '─'*100)
print('6. DESCRIPTION QUALITY')
print('─'*100)
m_empty_desc = sum(1 for tc in manual_tcs if not tc.get('description') or len(tc['description']) < 10)
a_empty_desc = sum(1 for tc in auto_tcs if not tc.get('description') or len(tc['description']) < 10)
m_avg_desc = sum(len(tc.get('description', '')) for tc in manual_tcs) / max(len(manual_tcs), 1)
a_avg_desc = sum(len(tc.get('description', '')) for tc in auto_tcs) / max(len(auto_tcs), 1)

print(f'    {"Metric":<35} {"Manual":>10} {"Auto":>10}')
print(f'    {"─"*35} {"─"*10} {"─"*10}')
print(f'    {"Empty/short descriptions":<35} {m_empty_desc:>10} {a_empty_desc:>10}')
print(f'    {"Avg description length (chars)":<35} {m_avg_desc:>10.0f} {a_avg_desc:>10.0f}')

# ─── SECTION 7: NAMING CONVENTION ANALYSIS ──────────────────────────────────
print('\n' + '─'*100)
print('7. NAMING CONVENTION ANALYSIS')
print('─'*100)

def analyze_naming(tcs_list, label):
    has_feature_id = sum(1 for tc in tcs_list if 'MWTGPROV' in tc['summary'])
    has_tc_prefix = sum(1 for tc in tcs_list if re.match(r'.*TC[-_]?\d+', tc['summary'], re.IGNORECASE))
    has_channel = sum(1 for tc in tcs_list if any(ch in tc['summary'].upper() for ch in ['ITMBO', 'NBOP', 'NSLNM']))
    has_api_ref = sum(1 for tc in tcs_list if any(a in tc['summary'].lower() for a in ['api', 'endpoint', 'rest', '/v1', '/v2']))
    avg_len = sum(len(tc['summary']) for tc in tcs_list) / max(len(tcs_list), 1)
    return {
        'Feature ID ref': has_feature_id,
        'TC## prefix': has_tc_prefix,
        'Channel ref': has_channel,
        'API ref': has_api_ref,
        'Avg name length': f'{avg_len:.0f}',
    }

m_naming = analyze_naming(manual_tcs, 'Manual')
a_naming = analyze_naming(auto_tcs, 'Auto')

print(f'    {"Metric":<35} {"Manual":>10} {"Auto":>10}')
print(f'    {"─"*35} {"─"*10} {"─"*10}')
for key in m_naming:
    print(f'    {key:<35} {str(m_naming[key]):>10} {str(a_naming[key]):>10}')

# ─── SECTION 8: DOCX CONTENT DUMP (for visibility) ──────────────────────────
print('\n' + '─'*100)
print('8. AUTO DOCX CONTENT PREVIEW (first 80 lines)')
print('─'*100)
for line in docx_text[:80]:
    if line:
        print(f'    {line[:120]}')

if auto_table_tcs:
    print(f'\n  AUTO DOCX TABLE DATA ({len(auto_table_tcs)} rows):')
    for i, row in enumerate(auto_table_tcs[:10]):
        print(f'    Row {i}: {row}')

# ─── SECTION 9: FULL TC LISTING ─────────────────────────────────────────────
print('\n' + '─'*100)
print('9. FULL TC LISTING')
print('─'*100)

print('\n  ── MANUAL TCs ──')
for tc in manual_tcs:
    print(f'    [{tc["sno"]:>3}] {tc["summary"]}')
    if tc.get('category'):
        print(f'          Category: {tc["category"]}')
    if tc.get('priority'):
        print(f'          Priority: {tc["priority"]}')

print('\n  ── AUTO TCs ──')
for tc in auto_tcs:
    print(f'    [{tc["sno"]:>3}] {tc["summary"]}')
    if tc.get('category'):
        print(f'          Category: {tc["category"]}')

# ─── SECTION 10: FINDINGS & ACTION ITEMS ────────────────────────────────────
print('\n' + '='*100)
print('10. FINDINGS & ACTION ITEMS')
print('='*100)

print(f"""
  ┌─────────────────────────────────────────────────────────────────────────────────┐
  │ SUMMARY                                                                         │
  │ Manual TCs: {len(manual_tcs):<5}  |  Auto TCs: {len(auto_tcs):<5}  |  Delta: {len(auto_tcs)-len(manual_tcs):>+4}                       │
  │ Manual Steps: {m_total_steps:<5}|  Auto Steps: {a_total_steps:<5}|  Avg Steps: {m_avg_steps:.1f} vs {a_avg_steps:.1f}          │
  └─────────────────────────────────────────────────────────────────────────────────┘

  Review the output above for detailed gap analysis.
  Key areas to check:
    1. Missing scenarios in auto-generated suite
    2. Naming convention alignment
    3. Step granularity differences
    4. Description completeness
    5. Category/type coverage balance
""")
